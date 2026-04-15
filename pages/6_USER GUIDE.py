import streamlit as st
from pathlib import Path
import io
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
import base64


def add_watermark():
    logo_path = Path("assets") / "fbc_logo.png"
    if logo_path.exists():
        with open(logo_path, "rb") as f:
            logo_base64 = base64.b64encode(f.read()).decode()

        watermark_css = f"""
        <style>

        /* Make watermark very light */
        .stApp::before {{
            content: "";
            position: fixed;
            top: 40;
            left: 50;
            width: 100%;
            height: 100%;
            background-image: url("data:image/png;base64,{logo_base64}");
            background-repeat: no-repeat;
            background-position: center;
            background-size: 1500px;
            opacity: 0.07;   /* 🔥 control watermark visibility here */
            pointer-events: none;
            z-index: 0;
        }}

        .block-container {{
            position: relative;
            z-index: 1;
        }}
        </style>
        """
        st.markdown(watermark_css, unsafe_allow_html=True)


add_watermark()
st.set_page_config(page_title="Help & Guide", layout="wide")

# ─── FBC DESIGN SYSTEM ─────────────────────────────────────────
st.markdown('''
<style>
/* ================================================================
   FBC INVESTMENT VALUATION SYSTEM — Design System v3.0
   ================================================================ */

/* ── 0. GOOGLE FONTS ──────────────────────────────────────── */
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700;900&family=EB+Garamond:ital,wght@0,400;0,600;1,400&family=Material+Icons&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined');

/* ── 1. GLOBAL TYPOGRAPHY ─────────────────────────────────── */
html, body, .stApp, .block-container,
p, div, label, span,
h1, h2, h3, h4, h5, h6,
li, ul, ol, a, small,
.stDataFrame, .stTable {
  font-family: "EB Garamond", Georgia, "Times New Roman", serif !important;
  color: #1a1a2e;
}

/* Headings — Playfair Display */
h1, h2, h3, h4, .fbc-heading, .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
  font-family: "Playfair Display", Georgia, serif !important;
  font-weight: 700 !important;
  letter-spacing: -0.01em !important;
}

/* ── 2. PAGE BACKGROUND ───────────────────────────────────── */
.stApp {
  background: #f5f7fb !important;
}
.main .block-container {
  background: #f5f7fb !important;
  padding-top: 1.5rem !important;
}

/* ── 3. SIDEBAR ───────────────────────────────────────────── */
section[data-testid="stSidebar"] {
    background: linear-gradient(175deg, #001a5c 0%, #003399 45%, #0044cc 100%) !important;
    border-right: 2px solid rgba(245,180,0,0.25) !important;
    box-shadow: 4px 0 24px rgba(0,26,92,0.35) !important;
}
section[data-testid="stSidebar"]::before {
    content: "";
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 3px;
    background: linear-gradient(90deg, #f5b400, #ffd040, #f5b400);
}
section[data-testid="stSidebar"] * {
    color: #e8f0ff !important;
    font-family: "EB Garamond", Georgia, serif !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
    font-family: "Playfair Display", serif !important;
    font-weight: 700 !important;
    text-shadow: 0 1px 4px rgba(0,0,0,0.3) !important;
}
section[data-testid="stSidebar"] .block-container {
    padding-top: 1.5rem !important;
}
section[data-testid="stSidebar"] a {
    color: #b8d0ff !important;
    transition: color 0.15s !important;
}
section[data-testid="stSidebar"] a:hover {
    color: #f5b400 !important;
}

/* sidebar hr divider */
section[data-testid="stSidebar"] hr {
    border: none !important;
    border-top: 1px solid rgba(245,180,0,0.25) !important;
    margin: 12px 0 !important;
}

/* ── 4. SIDEBAR COLLAPSE BUTTON ───────────────────────────── */
.material-icons,
span.material-icons,
i.material-icons,
.material-symbols-outlined,
[data-testid="stSidebarCollapseButton"] span,
[data-testid="stSidebarCollapseButton"] i {
    font-family: "Material Icons", "Material Symbols Outlined" !important;
    font-weight: normal !important;
    font-style: normal !important;
    letter-spacing: normal !important;
    text-transform: none !important;
    direction: ltr !important;
    -webkit-font-feature-settings: 'liga' !important;
    -webkit-font-smoothing: antialiased !important;
}
[data-testid="stSidebarCollapseButton"] button {
    background: linear-gradient(135deg, #f5b400, #ffd040) !important;
    border: none !important;
    border-radius: 50% !important;
    width: 44px !important; height: 44px !important;
    box-shadow: 0 4px 14px rgba(245,180,0,0.50) !important;
    transition: all 0.2s ease !important;
}
[data-testid="stSidebarCollapseButton"] button:hover {
    transform: translateY(-1px) scale(1.06) !important;
    box-shadow: 0 8px 20px rgba(245,180,0,0.60) !important;
}
[data-testid="stSidebarCollapseButton"] svg {
    width: 22px !important; height: 22px !important;
    fill: #001a5c !important;
}

/* ── 5. PAGE HEADER BANNER ────────────────────────────────── */
.fbc-page-header {
    background: linear-gradient(135deg, #001a5c 0%, #003399 50%, #0044cc 100%);
    border-radius: 18px;
    padding: 26px 32px;
    margin-bottom: 28px;
    border: 1px solid rgba(255,255,255,0.08);
    border-bottom: 3px solid #f5b400;
    box-shadow: 0 12px 40px rgba(0,26,92,0.28);
    position: relative;
    overflow: hidden;
}
.fbc-page-header::before {
    content: "";
    position: absolute;
    left: -40px; top: -40px;
    width: 200px; height: 200px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(245,180,0,0.12), transparent 70%);
    pointer-events: none;
}
.fbc-page-header::after {
    content: "";
    position: absolute;
    right: -30px; bottom: -30px;
    width: 160px; height: 160px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(255,255,255,0.06), transparent 65%);
    pointer-events: none;
}
.fbc-page-header-icon {
    font-size: 30px;
    margin-right: 12px;
    vertical-align: middle;
}
.fbc-page-header-title {
    font-family: "Playfair Display", serif !important;
    font-size: 28px !important;
    font-weight: 900 !important;
    color: #ffffff !important;
    display: inline !important;
    vertical-align: middle !important;
    letter-spacing: -0.01em !important;
}
.fbc-page-header-sub {
    font-size: 14px;
    color: rgba(255,255,255,0.78) !important;
    margin-top: 8px;
    font-style: italic;
}
.fbc-badge {
    display: inline-block;
    background: rgba(245,180,0,0.20);
    border: 1px solid rgba(245,180,0,0.50);
    color: #f5c842 !important;
    font-size: 10px;
    font-weight: 800;
    letter-spacing: 0.12em;
    padding: 3px 10px;
    border-radius: 999px;
    margin-left: 10px;
    vertical-align: middle;
    text-transform: uppercase;
    font-family: "EB Garamond", serif !important;
}

/* ── 6. SECTION HEADINGS ──────────────────────────────────── */
.fbc-section-heading {
    display: flex;
    align-items: center;
    gap: 12px;
    margin: 32px 0 16px 0;
    padding-bottom: 10px;
    border-bottom: 2px solid transparent;
    border-image: linear-gradient(90deg, #003399 0%, #f5b400 55%, transparent 90%) 1;
}
.fbc-section-heading-text {
    font-family: "Playfair Display", serif !important;
    font-size: 18px !important;
    font-weight: 700 !important;
    color: #001a5c !important;
    letter-spacing: 0.01em !important;
}
.fbc-section-heading-step {
    background: linear-gradient(135deg, #003399, #0044cc);
    color: white !important;
    font-size: 11px;
    font-weight: 900;
    min-width: 28px; height: 28px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    flex-shrink: 0;
    box-shadow: 0 3px 8px rgba(0,51,153,0.35);
}
.fbc-subsection-heading {
    font-family: "Playfair Display", serif !important;
    font-size: 15px !important;
    font-weight: 700 !important;
    color: #003399 !important;
    margin: 20px 0 10px 0 !important;
    padding-left: 12px !important;
    border-left: 3px solid #f5b400 !important;
}

/* ── 7. CARD / PANEL ──────────────────────────────────────── */
.fbc-card {
    background: #ffffff;
    border: 1px solid rgba(0,51,153,0.09);
    border-left: 5px solid #003399;
    border-radius: 16px;
    padding: 20px 24px;
    margin-bottom: 18px;
    box-shadow: 0 6px 18px rgba(0,26,92,0.07);
    transition: box-shadow 0.2s, transform 0.2s;
}
.fbc-card:hover {
    box-shadow: 0 12px 32px rgba(0,51,153,0.14);
    transform: translateY(-2px);
}
.fbc-card h3, .fbc-card h4 {
    font-family: "Playfair Display", serif !important;
    color: #001a5c !important;
    margin: 0 0 10px 0 !important;
}
.fbc-card-gold {
    border-left-color: #f5b400 !important;
}
.fbc-card-green {
    border-left-color: #10b981 !important;
}

/* sub-card (nested) */
.fbc-subcard {
    background: rgba(0,51,153,0.03);
    border: 1px solid rgba(0,51,153,0.10);
    border-radius: 12px;
    padding: 16px 18px;
    margin-top: 12px;
}

/* ── 8. KPI METRIC CARDS ──────────────────────────────────── */
.fbc-kpi {
    background: linear-gradient(135deg, #f0f5ff, #fff8e6);
    border: 1px solid rgba(0,51,153,0.12);
    border-radius: 16px;
    padding: 16px 18px;
    text-align: center;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    transition: transform 0.2s, box-shadow 0.2s;
}
.fbc-kpi:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 20px rgba(0,51,153,0.12);
}
.fbc-kpi-label {
    font-size: 10px;
    font-weight: 700;
    letter-spacing: 0.10em;
    text-transform: uppercase;
    color: #5a7099 !important;
    margin-bottom: 6px;
    font-family: "EB Garamond", serif !important;
}
.fbc-kpi-value {
    font-family: "Playfair Display", serif !important;
    font-size: 24px !important;
    font-weight: 800 !important;
    color: #001a5c !important;
    line-height: 1.2;
}
.fbc-kpi-unit {
    font-size: 11px;
    color: #7a90b8 !important;
    margin-top: 3px;
}

/* DCF card / kpi aliases */
.dcf-card {
    background: #ffffff;
    border: 1px solid rgba(0,0,0,0.08);
    border-left: 5px solid #003399;
    border-radius: 16px;
    padding: 20px 22px;
    box-shadow: 0 6px 18px rgba(0,0,0,0.06);
    margin-top: 12px;
    margin-bottom: 16px;
}
.dcf-card h3 { margin: 0 0 10px 0; font-family: "Playfair Display", serif !important; color: #001a5c !important; }
.dcf-subcard {
    background: rgba(0,51,153,0.03);
    border: 1px solid rgba(0,51,153,0.10);
    border-radius: 14px;
    padding: 14px;
    margin-top: 10px;
}
.dcf-kpi {
    background: linear-gradient(135deg, rgba(0,51,153,0.08), rgba(245,180,0,0.08));
    border: 1px solid rgba(0,0,0,0.08);
    border-radius: 16px;
    padding: 14px 16px;
    margin: 6px 0;
    transition: transform 0.2s;
}
.dcf-kpi:hover { transform: translateY(-1px); }
.dcf-kpi-title { font-size: 11px; opacity: 0.7; margin-bottom: 3px; text-transform: uppercase; letter-spacing: 0.07em; }
.dcf-kpi-value { font-size: 20px; font-weight: 800; font-family: "Playfair Display", serif !important; color: #001a5c !important; }

/* ── 9. RESET / UTILITY CARDS ─────────────────────────────── */
.fbc-reset-card {
    background: linear-gradient(135deg, #001a5c 0%, #003399 55%, #0055cc 100%);
    padding: 22px 28px;
    border-radius: 16px;
    color: white !important;
    box-shadow: 0 8px 24px rgba(0,26,92,0.30);
    margin-bottom: 22px;
    border-bottom: 3px solid #f5b400;
    position: relative;
    overflow: hidden;
}
.fbc-reset-card::after {
    content: "";
    position: absolute;
    right: -20px; top: -20px;
    width: 120px; height: 120px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(245,180,0,0.15), transparent 70%);
    pointer-events: none;
}
.fbc-reset-title {
    font-family: "Playfair Display", serif !important;
    font-size: 20px !important;
    font-weight: 700 !important;
    margin-bottom: 6px;
    color: white !important;
}
.fbc-reset-sub { font-size: 14px; opacity: 0.88; margin-bottom: 14px; color: rgba(255,255,255,0.85) !important; }

/* ── 10. BUTTONS ──────────────────────────────────────────── */
.stButton > button[kind="primary"],
.stButton > button {
    background: linear-gradient(135deg, #003399, #0044cc) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    font-family: "EB Garamond", Georgia, serif !important;
    padding: 10px 20px !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 12px rgba(0,51,153,0.28) !important;
    letter-spacing: 0.01em !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #0044cc, #0055ee) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 22px rgba(0,51,153,0.38) !important;
}
.fbc-reset-btn .stButton > button {
    background: linear-gradient(135deg, #f5b400, #ffd040) !important;
    color: #001a5c !important;
    box-shadow: 0 4px 12px rgba(245,180,0,0.32) !important;
}
.fbc-reset-btn .stButton > button:hover {
    background: linear-gradient(135deg, #ffd040, #ffe070) !important;
    box-shadow: 0 8px 20px rgba(245,180,0,0.45) !important;
}
.fbc-nav-btn button,
.fbc-nav-btn .stButton > button {
    background: linear-gradient(135deg, #003399, #001a4d) !important;
    color: white !important;
    font-weight: 700 !important;
    border-radius: 10px !important;
    padding: 8px 18px !important;
    border: none !important;
    transition: all 0.25s ease-in-out !important;
}
.fbc-nav-btn button:hover,
.fbc-nav-btn .stButton > button:hover {
    background: linear-gradient(135deg, #0055cc, #003399) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 14px rgba(0,0,0,0.25) !important;
}
.stDownloadButton > button {
    background: linear-gradient(135deg, #003399, #0044cc) !important;
    color: white !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    border: none !important;
    transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #0044cc, #0055ee) !important;
    transform: translateY(-1px) !important;
}

/* ── 11. INPUTS ───────────────────────────────────────────── */
.stNumberInput input, .stTextInput input {
    border: 1.5px solid rgba(0,51,153,0.18) !important;
    border-radius: 10px !important;
    font-family: "EB Garamond", Georgia, serif !important;
    background: #ffffff !important;
    transition: border-color 0.15s, box-shadow 0.15s !important;
    padding: 8px 12px !important;
}
.stNumberInput input:focus, .stTextInput input:focus {
    border-color: #003399 !important;
    box-shadow: 0 0 0 3px rgba(0,51,153,0.10) !important;
}
.stSelectbox > div > div {
    border: 1.5px solid rgba(0,51,153,0.18) !important;
    border-radius: 10px !important;
    font-family: "EB Garamond", Georgia, serif !important;
}

/* ── 12. TABS ─────────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px;
    border-bottom: 2px solid rgba(0,51,153,0.12);
    padding-bottom: 0;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 10px 10px 0 0 !important;
    font-weight: 700 !important;
    font-family: "EB Garamond", Georgia, serif !important;
    color: #5a7099 !important;
    padding: 10px 20px !important;
    transition: all 0.15s !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #003399, #0044cc) !important;
    color: white !important;
    box-shadow: 0 -2px 0 0 #f5b400 inset !important;
}

/* ── 13. EXPANDERS ────────────────────────────────────────── */
.streamlit-expanderHeader {
    font-weight: 700 !important;
    color: #001a5c !important;
    font-family: "Playfair Display", Georgia, serif !important;
    border-radius: 10px !important;
    background: rgba(0,51,153,0.04) !important;
    padding: 12px 16px !important;
    border-left: 3px solid #f5b400 !important;
}
/* Fix: hide raw "keyboard_arrow_right" text Streamlit injects into expander headers */
[data-testid="stExpander"] details > summary > div > p > span[style],
[data-testid="stExpander"] details > summary p span[data-testid="stMarkdownContainer"] > p > span,
details > summary p > span[style*="font-weight"] {
    display: none !important;
}
[data-testid="stExpanderToggleIcon"] svg,
details > summary svg {
    font-family: "Material Icons", "Material Symbols Outlined" !important;
}

/* ── 14. METRICS (st.metric) ──────────────────────────────── */
[data-testid="metric-container"] {
    background: linear-gradient(135deg, #f0f5ff, #fff8e6) !important;
    border: 1px solid rgba(0,51,153,0.12) !important;
    border-radius: 16px !important;
    padding: 16px 18px !important;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05) !important;
    transition: transform 0.2s, box-shadow 0.2s !important;
}
[data-testid="metric-container"]:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 20px rgba(0,51,153,0.10) !important;
}
[data-testid="metric-container"] label {
    font-size: 10px !important;
    font-weight: 700 !important;
    letter-spacing: 0.10em !important;
    text-transform: uppercase !important;
    color: #5a7099 !important;
    font-family: "EB Garamond", serif !important;
}
[data-testid="stMetricValue"] {
    font-family: "Playfair Display", serif !important;
    font-size: 24px !important;
    font-weight: 800 !important;
    color: #001a5c !important;
}
[data-testid="stMetricDelta"] {
    font-family: "EB Garamond", serif !important;
}

/* ── 15. ALERTS & INFO BOXES ──────────────────────────────── */
.stAlert {
    border-radius: 14px !important;
    font-family: "EB Garamond", Georgia, serif !important;
    border-left-width: 4px !important;
}

/* ── 16. PROGRESS BAR ─────────────────────────────────────── */
.stProgress > div > div > div {
    background: linear-gradient(90deg, #003399, #0044cc, #f5b400) !important;
    border-radius: 999px !important;
}

/* ── 17. SCROLLBAR ────────────────────────────────────────── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: #f0f4ff; border-radius: 999px; }
::-webkit-scrollbar-thumb { background: linear-gradient(#003399, #0044cc); border-radius: 999px; }
::-webkit-scrollbar-thumb:hover { background: #0055ee; }

/* ── 18. FOOTER ───────────────────────────────────────────── */
.fbc-footer {
    text-align: center;
    padding: 28px;
    margin-top: 48px;
    color: #5a7099 !important;
    font-size: 13px;
    border-top: 1px solid rgba(0,51,153,0.10);
    font-style: italic;
}
.fbc-footer b { color: #003399 !important; font-style: normal; }

/* ── 19. FEATURE CARDS (dashboard) ───────────────────────── */
.feature-box {
    background: #ffffff;
    padding: 22px 26px;
    border-radius: 16px;
    border-left: 6px solid #003399;
    box-shadow: 0 4px 14px rgba(0,0,0,0.07);
    transition: all 0.25s ease;
    margin-bottom: 16px;
    font-family: "EB Garamond", serif !important;
}
.feature-box:hover {
    background: linear-gradient(135deg, #f4f8ff, #fffdf0);
    box-shadow: 0 10px 28px rgba(0,51,153,0.16);
    transform: translateY(-3px);
    border-left-color: #f5b400;
}
.feature-icon { font-size: 24px; margin-right: 10px; }

/* ── 20. SMALL UTILITIES ──────────────────────────────────── */
.small-note {
    font-size: 12px;
    color: #7a90b8 !important;
    font-style: italic;
}
.gold-accent { color: #b87c00 !important; font-weight: 700; }
.blue-accent  { color: #003399 !important; font-weight: 700; }

/* ── 21. DATAFRAMES ───────────────────────────────────────── */
.stDataFrame thead th {
    background: linear-gradient(135deg, #001a5c, #003399) !important;
    color: white !important;
    font-weight: 700 !important;
    font-family: "Playfair Display", serif !important;
    letter-spacing: 0.02em !important;
}
.stDataFrame tbody tr:hover td {
    background: rgba(0,51,153,0.04) !important;
}
.stDataFrame tbody tr:nth-child(even) td {
    background: rgba(0,51,153,0.02) !important;
}

/* ── 22. TOP NAV (dashboard) ──────────────────────────────── */
.top-nav {
    position: fixed; top: 0; left: 0; width: 100%; height: 68px;
    background: linear-gradient(90deg, #001a5c, #003399, #0044cc);
    color: white; display: flex; align-items: center;
    padding: 0 28px; z-index: 99999;
    box-shadow: 0 4px 16px rgba(0,0,0,0.30);
    border-bottom: 2px solid #f5b400;
}
.top-title {
    font-family: "Playfair Display", serif !important;
    font-size: 26px; font-weight: 900; margin-left: 16px; letter-spacing: -0.01em;
    color: white;
}

/* ── 23. DIVIDER ──────────────────────────────────────────── */
.fbc-divider {
    height: 2px;
    background: linear-gradient(90deg, transparent, #f5b400, #003399, transparent);
    border: none;
    margin: 28px 0;
    border-radius: 999px;
}

/* ── 24. RADIO / CHECKBOX ─────────────────────────────────── */
.stRadio > label, .stCheckbox > label {
    font-family: "EB Garamond", Georgia, serif !important;
    color: #1a1a2e !important;
}

/* ── 25. SELECTBOX OPTIONS ────────────────────────────────── */
div[data-baseweb="select"] span {
    font-family: "EB Garamond", Georgia, serif !important;
}

/* ── 26. PEER PICKER (Comparables) ───────────────────────── */
.peer-picker-wrap {
    border: 1px solid rgba(0,51,153,0.14);
    border-radius: 18px;
    padding: 18px 18px 12px 18px;
    background: #f8fbff;
    box-shadow: 0 6px 20px rgba(0,26,92,0.07);
    margin-top: 12px;
    margin-bottom: 14px;
}
.peer-picker-head {
    font-family: "Playfair Display", serif !important;
    font-size: 17px;
    font-weight: 700;
    color: #001a5c;
    margin-bottom: 8px;
}
.peer-picker-sub {
    font-size: 13px;
    color: #475569;
    margin-bottom: 14px;
    font-style: italic;
}

/* ── 27. SENS TABLE OVERRIDES ─────────────────────────────── */
.sens-outer { border: 2px solid rgba(0,51,153,0.15) !important; border-radius: 16px !important; }
.sens-table { font-family: "EB Garamond", Georgia, serif !important; }
.sens-table thead th { background: #001a5c !important; }

/* ── 28. FOOTER BANNER ────────────────────────────────────── */
.footer {
    text-align: center;
    padding: 24px;
    margin-top: 40px;
    color: #5a7099 !important;
    font-size: 13px;
    border-top: 1px solid rgba(0,51,153,0.10);
    font-style: italic;
}
.footer b { color: #003399 !important; font-style: normal; }

</style>
''', unsafe_allow_html=True)

st.markdown('''
<div class="fbc-page-header">
    <span class="fbc-page-header-icon">🧭</span>
    <span class="fbc-page-header-title">User Guide & Documentation</span>
    <span class="fbc-badge">FBC Securities</span>
    <div class="fbc-page-header-sub">Step-by-step walkthrough of the FBC Valuation System.</div>
</div>
''', unsafe_allow_html=True)
# ────────────────────────────────────────────────────────────────

# ---------------------------------------------------------
# ✅ FIX SIDEBAR COLLAPSE ARROW (Material Icons)
# ---------------------------------------------------------
st.markdown("""
<style>

/* ✅ Load Material Icons so Streamlit's sidebar collapse icon renders correctly */
@import url('https://fonts.googleapis.com/icon?family=Material+Icons');

/* ✅ Make sure ONLY icons use the Material Icons font (prevents 'keyboard_double_arrow_right' text) */
.material-icons, 
span.material-icons,
i.material-icons,
[data-testid="stSidebarCollapseButton"] span,
[data-testid="stSidebarCollapseButton"] i {
    font-family: 'Material Icons' !important;
    font-weight: normal !important;
    font-style: normal !important;
    letter-spacing: normal !important;
    text-transform: none !important;
    display: inline-block !important;
    white-space: nowrap !important;
    word-wrap: normal !important;
    direction: ltr !important;
    -webkit-font-feature-settings: 'liga' !important;
    -webkit-font-smoothing: antialiased !important;
}

/* ✅ Style the collapse/expand button nicely */
[data-testid="stSidebarCollapseButton"] button {
    background: #003399 !important;
    border: 1px solid rgba(255,255,255,0.25) !important;
    border-radius: 999px !important;
    width: 44px !important;
    height: 44px !important;
    box-shadow: 0 6px 18px rgba(0, 51, 153, 0.35) !important;
    transition: transform 0.15s ease, box-shadow 0.15s ease, background 0.15s ease !important;
}

[data-testid="stSidebarCollapseButton"] button:hover {
    transform: translateY(-1px) !important;
    background: #0047d6 !important;
    box-shadow: 0 10px 22px rgba(0, 71, 214, 0.35) !important;
}

[data-testid="stSidebarCollapseButton"] svg {
    width: 22px !important;
    height: 22px !important;
    fill: white !important;
}
/* ===== SIDEBAR GLASS STYLE ===== */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #003399 0%, #001a4d 100%) !important;
    color: white !important;
    border-right: 1px solid rgba(255,255,255,0.15);
    backdrop-filter: blur(8px);
}

/* Sidebar text */
section[data-testid="stSidebar"] * {
    color: white !important;
}

/* Sidebar headings */
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
}

/* Remove default sidebar padding spacing */
section[data-testid="stSidebar"] .block-container {
    padding-top: 1rem !important;
}

</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# STYLES (GEORGIA FONT + BLUISH THEME)
# ---------------------------------------------------------
st.markdown(
    """
    <style>

      /* ===== GLOBAL FONT + COLOR ===== */
        html, body, .stApp, .block-container,
        p, div, label,
        h1, h2, h3, h4, h5, h6,
        li, ul, ol, a, small {
          font-family: Georgia, "Times New Roman", serif !important;
        }

      /* ===== TITLE ===== */
      .main-title {
        font-size: 2.0rem;
        font-weight: 800;
        margin-bottom: 0.2rem;
        color: #1e3a8a;
      }

      .subtle {
        color: #3b82f6;   /* lighter blue subtitle */
        margin-top: 0;
      }


      /* ===== CARD STYLE ===== */
      .card {
        border: 1px solid #dbeafe;
        border-radius: 14px;
        padding: 16px 18px;
        background: #f8fbff;
        box-shadow: 0 4px 18px rgba(30, 58, 138, 0.15);
        margin-bottom: 14px;
      }

      .card h3 {
        margin: 0 0 8px 0;
        font-size: 1.1rem;
        color: #1e40af;
      }

      /* ===== PILL LABEL ===== */
      .pill {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        font-size: 0.85rem;
        background: #e0ecff;
        border: 1px solid #bfdbfe;
        margin-right: 8px;
        color: #1e3a8a;
      }

      /* ===== CALLOUTS ===== */
      .callout {
        border-left: 5px solid #2563eb;
        background: #eff6ff;
        padding: 12px 14px;
        border-radius: 10px;
        margin: 10px 0;
        color: #1e3a8a;
      }

      .warn {
        border-left: 5px solid #f59e0b;
        background: #fffbeb;
        color: #92400e;
      }

      .danger {
        border-left: 5px solid #ef4444;
        background: #fef2f2;
        color: #991b1b;
      }

      /* ===== MONO TEXT (still Georgia as requested) ===== */
      .mono {
        font-family: Georgia, "Times New Roman", serif !important;
        font-size: 0.95rem;
        color: #1e3a8a;
      }

      hr {
        margin: 1rem 0;
        border: 1px solid #dbeafe;
      }

    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# HEADER
# ---------------------------------------------------------
st.markdown('<div class="main-title">🧭 Help & User Guide</div>', unsafe_allow_html=True)
st.markdown(
    "<p class='subtle'>Everything you need to use the valuation app smoothly — inputs, formulas, exports, and troubleshooting.</p>",
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# QUICK NAV / OVERVIEW
# ---------------------------------------------------------
colA, colB, colC = st.columns([1.2, 1, 1])

with colA:
    st.markdown(
        """
        <div class="card">
          <h3>✅ Quick Start</h3>
          <div class="callout">
            1) Open a valuation module (DCF / DDM / etc.)<br>
            2) Fill inputs step-by-step<br>
            3) Check results & sensitivity tables<br>
            4) Export Excel models for documentation
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with colB:
    st.markdown(
        """
        <div class="card">
          <h3>📌 What gets saved?</h3>
          <p class="subtle">
            The app uses <span class="mono">st.session_state</span> to keep your inputs
            across pages in the same session.
          </p>
          <div class="callout warn">
            Refreshing the browser may clear some values depending on your setup.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with colC:
    st.markdown(
        """
        <div class="card">
          <h3>🧾 Exports</h3>
          <p class="subtle">
            Each module can generate an Excel workbook with formulas, formatted sheets,
            and a summary page.
          </p>
          <div class="callout">
            Best practice: Export after finalizing your assumptions.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.markdown("---")

# ---------------------------------------------------------
# TABS PER MODULE
# ---------------------------------------------------------
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(
    ["⚡ Quick Summary", "💰 DDM", "📈 COMPARABLES", "🏦 BANKING", "🧾 SUMMARY", "🛠 Troubleshooting", "📉 DCF"])

# -------------------------
# DCF TAB
# -------------------------
with tab1:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">FAST</span> ⚡ Quick Summary — How to Value Using Each Model</h3>
          <p class="subtle">
            A 60-second guide. Follow these steps to get a valuation quickly (then use the other tabs for deeper detail).
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.markdown(
            """
            <div class="card">
              <h3>📉 DCF (UFCF / FCFF)</h3>
              <div class="callout">
                <b>Quick steps</b><br>
                1) Upload IS + BS + CF (one Excel, 3 sheets).<br></div>
                 <div style="margin-top:12px; padding:12px; border-radius:10px; background:#eef6ff; border-left:4px solid #2563eb;">
             <b>🔹 Required sheet order</b><br>
             • <b>Sheet 1:</b> Income Statement<br>
             • <b>Sheet 2:</b> Balance Sheet<br>
             • <b>Sheet 3:</b> Cash Flows
           </div>

           <div style="margin-top:10px; padding:12px; border-radius:10px; background:#f0f9ff; border-left:4px solid #1d4ed8;">
             <b>🔹 Required sheet layout</b><br>
             • <b>Column A</b> = line items (start on <b>Row 2</b>)<br>
             • <b>Row 1</b> (from <b>Column B</b> onward) = years (e.g., 2022, 2023, 2024)
           </div>

           <div style="margin-top:10px; padding:12px; border-radius:10px; background:#f8fafc; border-left:4px solid #475569;">
             <b>🔹 FX file (only if ZWG)</b><br>
             • Upload an FX Excel with <b>Date</b> + <b>Bank</b> and/or <b>Interbank</b> rate columns
           </div>
            </div>  <div class="callout">
             2) Select currency (USD or ZWG + FX file).<br>
             3) Map: Revenue, Debt, Cash, CA, CL, Equity, Capex, Depreciation (if available).<br>
             4) Choose forecast years + revenue growth method.<br>
             5) Confirm WC% method.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• Review the historical <b>Working Capital % of Sales</b> table.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• If any year looks abnormal, untick the <b>“Include”</b> box to exclude that year from the average calculation.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• Then choose whether to use:<br>
             &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;– The average of the included years, or<br>
             &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;– The most recent WC% for forecasting.<br>
             6) Enter <b>Average Cost of Debt Zimbabwe (US$) (%)</b>.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• Used to derive Risk-Free Rate (RF) when Auto mode is enabled.<br>
             7) Tick <b>“Use Auto (from Excel) for RF & MRP”</b> if using Country ERP + Default Spread file.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• RF and MRP populate automatically from Excel.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• You may untick to manually override RF and MRP.<br>
             8) Under <b>Select Industry / Industries (for blended βu)</b>:<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• Choose industries from the auto beta list,<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• OR manually override βu.<br>
             &nbsp;&nbsp;&nbsp;&nbsp;• Select simple or weighted average if multiple industries are chosen.<br>
             9) Enter Tax rate and <b>Terminal growth rate (g)</b>.<br>
             10) Select Valuation timing (valuation date and financial statement year-end date).<br>
             11) Review <b>CAPEX History — Exclude outlier years before averaging</b> if needed.<br>
             12) Review EV → Equity and the WACC vs g sensitivity grid.<br>
             13) Export Excel for audit trail. </div>

         </div>
         """,
            unsafe_allow_html=True,
        )

        st.markdown(
            """
            <div class="card">
              <h3>💰 DDM (Gordon Growth)</h3>
              <div class="callout">
                <b>Quick steps</b><br>
                1) Enter dividend history (prefer DPS if you want value per share).<br>
                2) Pick stable years for growth range (avoid special/irregular dividends).<br>
                3) Confirm computed g and D1.<br>
                4) Set CAPM inputs (RF, MRP, beta, D/E, tax) or override manually.<br>
                5) Check: <span class="mono">Ke &gt; g</span> then compute P0.<br>
                6) Enter shares to get total equity value.<br>
                7) Export Excel model.
              </div>
              <div class="callout warn">
                <b>Key rule:</b> If <span class="mono">Ke ≤ g</span>, the model is invalid.
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            """
            <div class="card">
              <h3>📈 Comparables (EV/EBITDA · P/B · P/E)</h3>

              <div class="callout">
                <b>Quick Setup</b><br>
                1) Gather all the necessary data on comparable companies and Ratios (EV/EBITDA, P/E, P/B).<br>
                3) Choose how many comparables you want to use.<br>
                2) Enter the peers and also multiples (EV/EBITDA, P/B, P/E).<br>
                3) Use Include flags to remove outliers (do not delete).<br>
                4) Enter Discount % → system computes implied multiples.
              </div>

              <div style="margin-top:12px; padding:12px; border-radius:10px; background:#eef6ff; border-left:4px solid #2563eb;">
                <b>🔹 Maintainable EBITDA</b><br>
                • Select EBITDA year range<br>
                • Enter weights (%) for each year<br>
                • Choose whether to apply timing (from DCF) or not
              </div>

              <div style="margin-top:10px; padding:12px; border-radius:10px; background:#f0f9ff; border-left:4px solid #1d4ed8;">
                <b>🔹 Maintainable Earnings</b><br>
                • <b>Auto-applied from Maintainable EBITDA ONLY</b><br>
                • Uses the <b>same years</b>, <b>same weights (%)</b>, and <b>same timing choice</b><br>
                • Your job here is to <b>review</b> and confirm the earnings output
              </div>

              <div style="margin-top:10px; padding:12px; border-radius:10px; background:#f8fafc; border-left:4px solid #475569;">
                <b>🔹 Book Equity & Net Debt</b><br>
                • <b>Book Equity</b> is automatically pulled from the <b>DCF Model</b> (if available).<br>
                • If DCF is not used, it can be taken from <b>BANKING</b> where applicable.<br>
                • <b>Net Debt</b> is automatically pulled from the <b>DCF model</b>.<br>
                • You may manually override these values if needed.
              </div>

              <div class="callout warn" style="margin-top:12px;">
                <b>Important:</b><br>
                EBITDA logic can automatically flow into Earnings.<br>
                Run DCF (and Banking if applicable) first so Book Equity, Net Debt, 
                EBITDA, and Earnings populate automatically.
              </div>

            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            """
            <div class="card">
              <h3>🏦 Banking (Residual Income)</h3>
              <div class="callout">
                <b>Quick steps</b><br>
                1) Upload IS + BS + SoCE and select correct sheets.<br>
                2) If ZWG: upload FX and confirm average vs closing conversion logic.<br>
                3) Map Total Equity rows on BS (and SoCE closing total if needed).<br>
                4) Choose the best earnings line (Normalized profit / PAT / Net profit).<br>
                5) Choose base year (Year 0) → confirm BV0 and Earnings0.<br>
                6) Set Ke via CAPM (auto or manual) + forecast years.<br>
                7) Enter growth assumptions (BV growth, earnings growth, terminal g).<br>
                8) Check: <span class="mono">Ke &gt; g</span> then compute equity value + export.
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown(
        """
        <div class="card">
          <h3>🧾 Summary (Blended / Weighted Valuation)</h3>
          <div class="callout">
            <b>Quick steps</b><br>
            1) Run the valuation tabs you want (DCF/DDM/Comps/Banking) first.<br>
            2) In Summary, select models to include.<br>
            3) Input weights (the app normalizes them to 100%).<br>
            4) Enter shares and current share price.<br>
            5) Review intrinsic value, upside/downside, and recommendation.<br>
            6) Export Summary Excel.
          </div>
          <div class="callout warn">
            <b>Common issue:</b> If a model shows blank, it wasn’t run yet.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
# -------------------------
# DDM TAB
# -------------------------
with tab2:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">DDM</span> Dividend Discount Model (Gordon Growth)</h3>
          <p class="subtle">
            This module values equity by converting a growing stream of dividends into a single present value today.
            It’s best for stable, dividend-paying firms with predictable payout and long-run growth.
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if "guide_exp_17_1__what_the_ddm_page_does__big_picture_" not in st.session_state:
        st.session_state["guide_exp_17_1__what_the_ddm_page_does__big_picture_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_17_1__what_the_ddm_page_does__big_picture_"] else "▶ ") + "1) What the DDM page does (big picture)",
        key="btn_guide_exp_17_1__what_the_ddm_page_does__big_picture_"
    ):
        st.session_state["guide_exp_17_1__what_the_ddm_page_does__big_picture_"] = not st.session_state["guide_exp_17_1__what_the_ddm_page_does__big_picture_"]
    if st.session_state["guide_exp_17_1__what_the_ddm_page_does__big_picture_"]:
        st.markdown(r"""
**This DDM page produces an Equity Value per Share using the Gordon Growth Dividend Discount Model.**

It takes:
- **Dividend history** (you enter dividends per year),
- chooses a **growth range** (years used to compute dividend CAGR),
- calculates **growth rate (g)** and **next dividend (D₁)**,
- computes **Cost of Equity (Rₑ)** using CAPM (and **levered beta** using D/E and tax),
- outputs **Value per share (P₀)** and optionally **Total Equity Value** using shares outstanding,
- generates an **auditable Excel model** with steps + formulas.

### Core Formula
""")
        st.latex(r"P_0 = \frac{D_1}{R_e - g}")
        st.markdown(r"""
Where:
- **D₁** = Dividend expected next year  
- **Rₑ** = Cost of Equity (usually CAPM)  
- **g** = Long-term dividend growth rate  

✅ **Key rule:** The model only works when **Rₑ > g**.
""")

    if "guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_" not in st.session_state:
        st.session_state["guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"] else "▶ ") + "2) Step 1 — Dividend History (what to enter and why)",
        key="btn_guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"
    ):
        st.session_state["guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"] = not st.session_state["guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"]
    if st.session_state["guide_exp_18_2__step_1___dividend_history__what_to_enter_and_why_"]:
        st.markdown(r"""
### What you do on the page
1) Choose **Start Year** and **End Year**  
2) Enter the **Dividend for each year** (the app stores it so it doesn’t reset within the session)  
3) The page shows a table of Year vs Dividend

### What the model needs from this step
- A **clean dividend series** to estimate growth.
- The dividend in the **final selected year** will later become the base for **D₁**.

### Tips (so your valuation makes sense)
- Use **dividend per share (DPS)** if you want **P₀ per share**.
- Keep the dividends consistent (same units for all years).
- If dividends are irregular, consider choosing a smaller growth range (Step 2) focusing on stable years.
""")

    if "guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_" not in st.session_state:
        st.session_state["guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"] else "▶ ") + "3) Step 2 — Growth Calculation Range (select stable years)",
        key="btn_guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"
    ):
        st.session_state["guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"] = not st.session_state["guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"]
    if st.session_state["guide_exp_19_3__step_2___growth_calculation_range__select_stable_years_"]:
        st.markdown(r"""
### What you do
You select:
- **Growth start year**
- **Growth end year**

These years tell the model which part of history to use for growth.

### Why this matters
Dividend growth can be distorted by:
- special dividends,
- payout policy changes,
- one-off shocks.

So you should pick a range that reflects **“normal” long-term dividend behavior**.
""")

    if "guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__" not in st.session_state:
        st.session_state["guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"] else "▶ ") + "4) Step 3 — Dividend Growth (g) and next dividend (D₁)",
        key="btn_guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"
    ):
        st.session_state["guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"] = not st.session_state["guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"]
    if st.session_state["guide_exp_20_4__step_3___dividend_growth__g__and_next_dividend__d__"]:
        st.markdown(r"""
### How growth (g) is calculated
- If **start year = end year**, growth is **0%**.
- If dividends are positive, the model uses **CAGR**:

""")
        st.latex(r"g = \left(\frac{D_{end}}{D_{start}}\right)^{\frac{1}{(end-start)}} - 1")
        st.markdown(r"""
- If the starting dividend is 0 (or unusable), the model uses a **fallback** (e.g., 2%) to avoid breaking.

### How next dividend (D₁) is calculated
""")
        st.latex(r"D_1 = D_{end}\times(1+g)")
        st.markdown(r"""
✅ Interpretation:
- **g** is your long-run dividend growth assumption implied by the selected history.
- **D₁** is the dividend the model expects **next year**.

**Best practice:** For mature firms, g should usually be conservative and close to long-run economic growth.
""")

    if "guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values" not in st.session_state:
        st.session_state["guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"] else "▶ ") + "5) Step 4 — Cost of Equity (Rₑ) via CAPM (and why DCF values appear here)",
        key="btn_guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"
    ):
        st.session_state["guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"] = not st.session_state["guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"]
    if st.session_state["guide_exp_21_5__step_4___cost_of_equity__r___via_capm__and_why_dcf_values"]:
        st.markdown(r"""
### Where inputs come from
This DDM page tries to **reuse your DCF assumptions** if they exist in `st.session_state`, such as:
- **Risk-free rate (RF)**
- **Equity risk premium / Market risk premium (MRP)**
- **Tax rate**
- **Unlevered beta (βu)**
- **Debt/Equity (D/E)**

You can also tick **“Manually override parameters”** to enter custom values.

### Levered beta
The model converts **unlevered beta** to **levered beta** using D/E and tax:

""")
        st.latex(r"\beta_L = \beta_u \times \left(1 + (1 - Tax)\times\frac{D}{E}\right)")
        st.markdown(r"""
### CAPM Cost of Equity
""")
        st.latex(r"R_e = RF + \beta_L \times MRP")
        st.markdown(r"""
✅ Interpretation:
- Higher **β** or **MRP** increases **Rₑ** → lowers valuation.
- Higher **D/E** increases βL (financial risk) → increases **Rₑ**.

**Tip:** If your D/E is extreme or equity is near zero, βL may become huge — your valuation will become unrealistic.
""")

    if "guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val" not in st.session_state:
        st.session_state["guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"] else "▶ ") + "6) Step 5 — Equity Value per Share (P₀) and the critical validity check",
        key="btn_guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"
    ):
        st.session_state["guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"] = not st.session_state["guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"]
    if st.session_state["guide_exp_22_6__step_5___equity_value_per_share__p___and_the_critical_val"]:
        st.markdown(r"""
### What the model does
- It checks the Gordon Growth rule: **Rₑ must be greater than g**.
- If **Rₑ ≤ g**, the model stops the valuation and shows an error (because the denominator becomes zero or negative).

### If the rule passes, the model computes:
""")
        st.latex(r"P_0 = \frac{D_1}{R_e - g}")
        st.markdown(r"""
✅ Interpretation:
- Higher **D₁** increases P₀.
- Higher **Rₑ** decreases P₀.
- Higher **g** increases P₀ (but too high g can break the model).

<div class="callout warn">
  <b>Common DDM issue:</b> If <span class="mono">Rₑ ≤ g</span>, Gordon Growth breaks.<br>
  Fix by using a more conservative <b>g</b> or revisiting your <b>Rₑ</b> assumptions (beta, MRP, D/E, tax).
</div>
""", unsafe_allow_html=True)

    if "guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_" not in st.session_state:
        st.session_state["guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"] else "▶ ") + "7) Step 6 — Total Equity Value (P₀ × shares outstanding)",
        key="btn_guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"
    ):
        st.session_state["guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"] = not st.session_state["guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"]
    if st.session_state["guide_exp_23_7__step_6___total_equity_value__p____shares_outstanding_"]:
        st.markdown(r"""
### What you do
Enter **Number of Shares**.

### What the model computes
If shares > 0 and P₀ is valid:
""")
        st.latex(r"\text{Total Equity Value} = P_0 \times \text{Shares Outstanding}")
        st.markdown(r"""
✅ Interpretation:
- If you already have **per-share dividends (DPS)**, then P₀ is **per share**, so multiplying by shares gives total equity value.

**Tip:** Make sure the “dividends” you entered are truly per share (or else your total value will be off by a scale factor).
""")

    if "guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_" not in st.session_state:
        st.session_state["guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"] else "▶ ") + "8) Excel Export (FULL DDM model: sheets + what to check)",
        key="btn_guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"
    ):
        st.session_state["guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"] = not st.session_state["guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"]
    if st.session_state["guide_exp_24_8__excel_export__full_ddm_model__sheets___what_to_check_"]:
        st.markdown(r"""
When you click **Generate / Update FULL DDM Excel Model**, the app creates an auditable workbook with:

- **DividendHistory**: Year & Dividend table  
- **Growth**: growth range + formulas for D_start, D_end, g, and D₁  
- **Parameters**: CAPM inputs + βL + Rₑ formulas  
- **Valuation**: Gordon Growth valuation + shares + total equity value  
- **Summary**: key outputs in one view

### Why this export matters
- It preserves formulas (INDEX/MATCH, CAPM, Gordon Growth).
- It is perfect for reporting, audit trail, and sharing assumptions.
""")

    if "guide_exp_25_9__troubleshooting__ddm_specific_" not in st.session_state:
        st.session_state["guide_exp_25_9__troubleshooting__ddm_specific_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_25_9__troubleshooting__ddm_specific_"] else "▶ ") + "9) Troubleshooting (DDM-specific)",
        key="btn_guide_exp_25_9__troubleshooting__ddm_specific_"
    ):
        st.session_state["guide_exp_25_9__troubleshooting__ddm_specific_"] = not st.session_state["guide_exp_25_9__troubleshooting__ddm_specific_"]
    if st.session_state["guide_exp_25_9__troubleshooting__ddm_specific_"]:
        st.markdown(r"""
### “Start year cannot be greater than end year”
- Your dividend history year range is invalid. Set Start Year ≤ End Year.

### “Growth start year must be earlier or equal to end year”
- Your growth range selection is reversed. Choose a valid range.

### “Rₑ must be greater than g”
- Gordon Growth is invalid when Rₑ ≤ g.
- Fix by:
  - lowering **g** (use a conservative perpetual growth),
  - checking **beta**, **MRP**, **D/E**, **tax**, **RF**.

### Excel download button disabled
- The download button is disabled until you click **Generate / Update**.
- Click generate first. If it still fails, check terminal logs for openpyxl errors.
""")

# -------------------------
# COMPARABLES TAB
# -------------------------
with tab3:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">COMPS</span> Comparable Companies Valuation</h3>
          <p class="subtle">
            This module values a target company using market trading multiples from comparable (“peer”) companies.
            It supports <b>EV/EBITDA</b>, <b>P/B</b>, and <b>P/E</b> methods, with optional peer-universe auto-fill and
            an Excel export that contains the full audit trail (inputs + formulas).
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if "guide_exp_26_1__what_this_comparables_page_does__big_picture_" not in st.session_state:
        st.session_state["guide_exp_26_1__what_this_comparables_page_does__big_picture_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_26_1__what_this_comparables_page_does__big_picture_"] else "▶ ") + "1) What this Comparables page does (big picture)",
        key="btn_guide_exp_26_1__what_this_comparables_page_does__big_picture_"
    ):
        st.session_state["guide_exp_26_1__what_this_comparables_page_does__big_picture_"] = not st.session_state["guide_exp_26_1__what_this_comparables_page_does__big_picture_"]
    if st.session_state["guide_exp_26_1__what_this_comparables_page_does__big_picture_"]:
        st.markdown(
            r"""
**Comparable Company Analysis (CCA)** estimates the target’s equity value by applying **peer trading multiples**
to the target’s own financial base (EBITDA, earnings, book equity).

This page does 6 main things:

1) **Auto-fills peers** from a Peer Universe Excel (optional but recommended)  
2) Lets you **review/edit comparables** and choose which multiples to include  
3) Computes **Average multiples** and applies a **Discount factor** to get “Implied” multiples  
4) Builds **Maintainable EBITDA** (from DCF) with weights (and optional timing effect)  
5) Builds **Maintainable Earnings** (from DCF) with weights (and optional timing effect)  
6) Computes **Equity values** using:
   - **EV/EBITDA** → Equity = (Implied EV/EBITDA × Maintainable EBITDA) − Net Debt  
   - **P/B** → Equity = (Implied P/B × Book Equity)  
   - **P/E** → Equity = (Implied P/E × Maintainable Earnings)  

✅ Outputs are displayed on-screen and can be exported to Excel with formulas.
"""
        )

    if "guide_exp_27_2__auto_fill_comparables__peer_universe_" not in st.session_state:
        st.session_state["guide_exp_27_2__auto_fill_comparables__peer_universe_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_27_2__auto_fill_comparables__peer_universe_"] else "▶ ") + "2) Auto-Fill Comparables (Peer Universe)",
        key="btn_guide_exp_27_2__auto_fill_comparables__peer_universe_"
    ):
        st.session_state["guide_exp_27_2__auto_fill_comparables__peer_universe_"] = not st.session_state["guide_exp_27_2__auto_fill_comparables__peer_universe_"]
    if st.session_state["guide_exp_27_2__auto_fill_comparables__peer_universe_"]:
        st.markdown(
            r"""
### Purpose
Step 0 helps you build a clean peer set quickly from a **Peer Universe file** (peer_universe.xlsx).

### What you do
1) Turn ON: **“Use Peer Universe Excel to auto-fill comparables”**  
2) Choose a **Target Company** (the firm you are valuing)  
3) Set **Max peers** (how many peers to suggest)  
4) Decide if you want auto-fill to happen instantly:
   - ✅ **Auto-fill comparables instantly when I choose a target**
5) Review the auto-selected peers in the **Peer companies** multiselect (you can edit it)

### How peers are suggested (important logic)
- If the target has a **PeerGroup**, the system uses **ONLY that PeerGroup** (strict matching).
  This prevents cross-industry mistakes.
- If PeerGroup is empty, it falls back to:
  **Industry → Sector** (and picks peers with more available multiples first).

### Uploading your own peer universe
- You can upload another Excel file using the uploader.
- The app stores it in memory for the session, so it stays even if you switch tabs.

### “Clear Comparables”
This button resets:
- company names,
- multiples,
- include/exclude flags,
so you can start fresh.
"""
        )

    if "guide_exp_28_3__step_1___input_comparable_companies___multiples" not in st.session_state:
        st.session_state["guide_exp_28_3__step_1___input_comparable_companies___multiples"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_28_3__step_1___input_comparable_companies___multiples"] else "▶ ") + "3) Step 1 — Input Comparable Companies & Multiples",
        key="btn_guide_exp_28_3__step_1___input_comparable_companies___multiples"
    ):
        st.session_state["guide_exp_28_3__step_1___input_comparable_companies___multiples"] = not st.session_state["guide_exp_28_3__step_1___input_comparable_companies___multiples"]
    if st.session_state["guide_exp_28_3__step_1___input_comparable_companies___multiples"]:
        st.markdown(
            r"""
### Purpose
This is where you confirm the peer set and enter (or review) their trading multiples.

### What you do
1) Set **How many comparables?**  
2) For each comparable, fill:
   - Company name
   - **EV/EBITDA**
   - **P/B**
   - **P/E**
3) Use the **Analyst filter** checkboxes to control which multiples are included in the averaging:
   - Include EV
   - Include P/B
   - Include P/E

### How the “Include” filters work
- If you uncheck “Include EV” for a company, that company’s EV/EBITDA is excluded from the EV averaging.
- The same idea applies for P/B and P/E.
- This lets you remove outliers or irrelevant peers *without deleting them*.

### Output you get
At the bottom you see a table of:
- Company
- Multiples
- Include flags
"""
        )

    if "guide_exp_29_4__step_2___average___implied_multiples__with_discount_" not in st.session_state:
        st.session_state["guide_exp_29_4__step_2___average___implied_multiples__with_discount_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_29_4__step_2___average___implied_multiples__with_discount_"] else "▶ ") + "4) Step 2 — Average & Implied Multiples (with discount)",
        key="btn_guide_exp_29_4__step_2___average___implied_multiples__with_discount_"
    ):
        st.session_state["guide_exp_29_4__step_2___average___implied_multiples__with_discount_"] = not st.session_state["guide_exp_29_4__step_2___average___implied_multiples__with_discount_"]
    if st.session_state["guide_exp_29_4__step_2___average___implied_multiples__with_discount_"]:
        st.markdown(
            r"""
### Purpose
Convert peer multiples into:
- a **simple average**, and then
- an **implied multiple** after applying a discount.

### What you do
1) Enter **Discount factor (%)**  
   Example: 25% means you reduce the peer multiple by 25%.

### What the model computes
For each multiple:

**Average multiple**  
- Uses only peers with the Include flag = True
- Uses the mean of the included values

**Implied multiple**
"""
        )
        st.latex(r"\text{Implied Multiple} = \text{Average Multiple} \times (1-\text{Discount})")
        st.markdown(
            r"""
✅ Interpretation:
- Higher discount → lower implied multiple → lower equity value.
- Discount is typically used to reflect:
  - size/illiquidity discount,
  - country risk,
  - control vs minority differences,
  - quality differences vs peers.

You will see a summary table with Average, Discount, and Implied for EV/EBITDA, P/B, and P/E.
"""
        )

    if "guide_exp_30_5__timing_source__from_dcf____why_it_exists" not in st.session_state:
        st.session_state["guide_exp_30_5__timing_source__from_dcf____why_it_exists"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_30_5__timing_source__from_dcf____why_it_exists"] else "▶ ") + "5) Timing Source (from DCF) — why it exists",
        key="btn_guide_exp_30_5__timing_source__from_dcf____why_it_exists"
    ):
        st.session_state["guide_exp_30_5__timing_source__from_dcf____why_it_exists"] = not st.session_state["guide_exp_30_5__timing_source__from_dcf____why_it_exists"]
    if st.session_state["guide_exp_30_5__timing_source__from_dcf____why_it_exists"]:
        st.markdown(
            r"""
### Purpose
This section pulls **DCF discount timing values (n)** to create a timing base used in:
- Maintainable EBITDA (Step 3)
- Maintainable Earnings (Step 4)

### What you do
- If DCF timing exists:
  - you can choose **Use DCF n₀** (recommended), or
  - manually override the starting timing value.
- If DCF timing does NOT exist:
  - you must enter a **manual timing base**.

### What “timing effect” means here
When timing is ON, each year gets a factor:
- Year 1 uses base_timing
- Year 2 uses base_timing + 1
- Year 3 uses base_timing + 2
…and so on.

This factor is applied before weighting to compute maintainable values.
"""
        )

    if "guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_" not in st.session_state:
        st.session_state["guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"] else "▶ ") + "6) Step 3 — Maintainable EBITDA (EV/EBITDA base)",
        key="btn_guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"
    ):
        st.session_state["guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"] = not st.session_state["guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"]
    if st.session_state["guide_exp_31_6__step_3___maintainable_ebitda__ev_ebitda_base_"]:
        st.markdown(
            r"""
### Purpose
Build a single “Maintainable EBITDA” value from DCF forecast EBITDA.

### Where EBITDA comes from
This module reads EBITDA from your DCF page (session_state keys like):
- dcf_ebitda_forecast / dcf_ebitda_all

If DCF EBITDA is missing:
- The EV/EBITDA method is skipped.

### What you do
1) Choose whether to apply timing:
   - **Apply timing effect from DCF to EBITDA?**
2) Select the EBITDA year range:
   - EBITDA Start Year
   - EBITDA End Year
3) Provide a **weight (%)** for each year in the selected range.

### What the model computes
For each year:
- **Timing factor** = 1 (if timing OFF) OR base_timing + index (if timing ON)
- **Adjusted EBITDA** = EBITDA × Timing
- **Weighted EBITDA** = Adjusted EBITDA × Weight

And then:
"""
        )
        st.latex(r"\text{Maintainable EBITDA} = \sum(\text{Weighted EBITDA})")
        st.markdown(
            r"""
✅ Tip (very important):
- Your weights are percentages. Make sure they make sense (many analysts aim for ~100% total,
  but the tool will still compute even if totals are above/below 100).

### Output you get
- A table showing EBITDA, timing, weights, adjusted EBITDA, and weighted EBITDA
- A final “Maintainable EBITDA” total
"""
        )

    if "guide_exp_32_7__step_4___maintainable_earnings__p_e_base_" not in st.session_state:
        st.session_state["guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"] else "▶ ") + "7) Step 4 — Maintainable Earnings (P/E base)",
        key="btn_guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"
    ):
        st.session_state["guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"] = not st.session_state["guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"]
    if st.session_state["guide_exp_32_7__step_4___maintainable_earnings__p_e_base_"]:
        st.markdown(
            r"""
### Purpose
Build a single “Maintainable Earnings” value from DCF forecast earnings.

### Where Earnings comes from
This module reads earnings from your DCF page (session_state keys like):
- dcf_profit_forecast / dcf_profit_all

If DCF Earnings is missing:
- The P/E method is skipped.

### Auto-sync features (important)
This page can keep Earnings weighting consistent with EBITDA weighting:

1) **Auto-use the SAME years & weights as EBITDA (recommended)**  
   - If ON: Earnings uses the same year range and weights as Step 3.  
   - This ensures method consistency.

2) **Timing is locked to EBITDA timing**  
   - If timing is OFF for EBITDA, timing is forced OFF for Earnings.

### What you do
1) Choose whether to sync years & weights to EBITDA  
2) Choose whether to apply timing (if allowed)  
3) If not syncing, you can manually choose the year range and weights.

### What the model computes
For each year:
- **Adjusted Earnings** = Earnings × Timing (or ×1 if timing off)
- **Weighted Earnings** = Adjusted Earnings × Weight

And then:
"""
        )
        st.latex(r"\text{Maintainable Earnings} = \sum(\text{Weighted Earnings})")
        st.markdown(
            r"""
### Output you get
- A table showing Earnings, timing, weights, adjusted and weighted values
- A final “Maintainable Earnings” total
"""
        )

    if "guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_" not in st.session_state:
        st.session_state["guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"] else "▶ ") + "8) Step 5 — Book Value & Net Debt (P/B and EV bridge)",
        key="btn_guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"
    ):
        st.session_state["guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"] = not st.session_state["guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"]
    if st.session_state["guide_exp_33_8__step_5___book_value___net_debt__p_b_and_ev_bridge_"]:
        st.markdown(
            r"""
### Purpose
Provide the balance sheet inputs needed for:
- P/B valuation (Book Equity)
- EV/EBITDA bridge (Net Debt)

### Book Equity auto-fill
If you used the Banking page, the tool can auto-fill Beginning Book Equity from:
- bank.outputs.book_equity_0

### What you do
1) Enter or confirm **Book Equity (USD)**  
2) Enter **Net Debt (USD)**  

✅ Interpretation:
- **Net Debt** is used to move from enterprise value to equity value in the EV/EBITDA method:
  Equity = Enterprise Value − Net Debt
"""
        )

    if "guide_exp_34_9__step_6___computed_equity_values__final_outputs_" not in st.session_state:
        st.session_state["guide_exp_34_9__step_6___computed_equity_values__final_outputs_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_34_9__step_6___computed_equity_values__final_outputs_"] else "▶ ") + "9) Step 6 — Computed Equity Values (final outputs)",
        key="btn_guide_exp_34_9__step_6___computed_equity_values__final_outputs_"
    ):
        st.session_state["guide_exp_34_9__step_6___computed_equity_values__final_outputs_"] = not st.session_state["guide_exp_34_9__step_6___computed_equity_values__final_outputs_"]
    if st.session_state["guide_exp_34_9__step_6___computed_equity_values__final_outputs_"]:
        st.markdown(
            r"""
### Purpose
Compute equity value using each comparable multiple method.

### Methods and formulas

**A) EV/EBITDA**
"""
        )
        st.latex(
            r"\text{Equity Value}_{EV/EBITDA} = (\text{Implied EV/EBITDA} \times \text{Maintainable EBITDA}) - \text{Net Debt}")
        st.markdown(
            r"""
**B) P/B**
"""
        )
        st.latex(r"\text{Equity Value}_{P/B} = (\text{Implied P/B} \times \text{Book Equity})")
        st.markdown(
            r"""
**C) P/E**
"""
        )
        st.latex(r"\text{Equity Value}_{P/E} = (\text{Implied P/E} \times \text{Maintainable Earnings})")
        st.markdown(
            r"""
### Output you get
A results table with:
- EV/EBITDA equity value
- P/B equity value
- P/E equity value
"""
        )

    if "guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu" not in st.session_state:
        st.session_state["guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"] else "▶ ") + "10) Excel Export — Comparables workbook (audit trail + formulas)",
        key="btn_guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"
    ):
        st.session_state["guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"] = not st.session_state["guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"]
    if st.session_state["guide_exp_35_10__excel_export___comparables_workbook__audit_trail___formu"]:
        st.markdown(
            r"""
### Purpose
Download an Excel file that reproduces the model with formulas.

### What the export includes
The exported workbook contains these sheets:

1) **Comps_Input**
   - peer names, multiples, include flags

2) **Multiples**
   - AVERAGEIF formulas using include flags
   - Discount input and implied multiples

3) **EBITDA_Maintainable**
   - timing toggle, base timing, EBITDA, weights, adjusted and weighted EBITDA
   - maintainable EBITDA total

4) **Earnings_Maintainable**
   - timing toggle, base timing, earnings, weights, adjusted and weighted earnings
   - maintainable earnings total

5) **Equity_Values**
   - book equity, net debt
   - links to implied multiples and maintainables
   - final equity values for EV/EBITDA, P/B, P/E

✅ Why this matters:
- It creates a clear audit trail for reporting and review.
- You can share it with stakeholders who prefer Excel.
"""
        )

    if "guide_exp_36_11__troubleshooting__common_issues_" not in st.session_state:
        st.session_state["guide_exp_36_11__troubleshooting__common_issues_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_36_11__troubleshooting__common_issues_"] else "▶ ") + "11) Troubleshooting (common issues)",
        key="btn_guide_exp_36_11__troubleshooting__common_issues_"
    ):
        st.session_state["guide_exp_36_11__troubleshooting__common_issues_"] = not st.session_state["guide_exp_36_11__troubleshooting__common_issues_"]
    if st.session_state["guide_exp_36_11__troubleshooting__common_issues_"]:
        st.markdown(
            r"""
### “Missing peer_universe.xlsx”
- Ensure **peer_universe.xlsx** is inside your project **/data/** folder,
  OR upload it using the uploader.

### “No timing values detected from DCF”
- Run the DCF model first (so timing exists), or manually set a timing base.

### “No EBITDA found from DCF — skipping EV/EBITDA method”
- Your DCF page has not populated EBITDA into session_state.
- Run DCF first or confirm the EBITDA keys are being stored correctly.

### “No Earnings found from DCF — skipping P/E method”
- Same idea: run DCF first or confirm earnings are stored correctly.

### Weird results from averages
- Check the **Include** flags.
- Check for outliers (e.g., one peer with a huge multiple).
- Ensure you didn’t unintentionally leave a peer with 0.00 that you meant to exclude.

### Excel values don’t match the screen
- Regenerate the Excel export after changing inputs (because Excel is created from current session_state).
"""
        )
# -------------------------
# BANKING TAB
# -------------------------
with tab4:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">BANKING</span> Banking Valuation — Residual Income Method</h3>
          <p class="subtle">
            This module values a bank (or financial institution) using the <b>Residual Income (RI)</b> approach
            with <b>actual year columns</b> from uploaded statements. It also supports optional <b>ZWG → USD</b>
            conversion using an FX Excel file (DCF-style).
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if "guide_exp_37_1__what_this_banking_page_does__big_picture_" not in st.session_state:
        st.session_state["guide_exp_37_1__what_this_banking_page_does__big_picture_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_37_1__what_this_banking_page_does__big_picture_"] else "▶ ") + "1) What this Banking page does (big picture)",
        key="btn_guide_exp_37_1__what_this_banking_page_does__big_picture_"
    ):
        st.session_state["guide_exp_37_1__what_this_banking_page_does__big_picture_"] = not st.session_state["guide_exp_37_1__what_this_banking_page_does__big_picture_"]
    if st.session_state["guide_exp_37_1__what_this_banking_page_does__big_picture_"]:
        st.markdown(
            r"""
### Residual Income Model (Banks)
For banks, traditional FCFF DCF can be tricky because debt is part of operating structure.  
Residual Income (RI) focuses on:
- **Book Value of Equity**, and
- **Earnings in excess of the equity charge (Ke × Book Value)**

**Core idea:**
- If a bank earns exactly Ke on book value → residual income = 0 → value ≈ book value
- If it earns more than Ke → positive residual income → value > book value

This page does 7 main things:
1) Upload your statements (IS + BS + SoCE)
2) (Optional) Convert ZWG → USD using FX Excel (yearly averages + BS closing rates)
3) Map the correct equity rows (Balance Sheet) and SoCE totals (for reference)
4) Select the best earnings line from the Income Statement
5) Choose a base year and pull base equity + base earnings
6) Compute Ke using a CAPM block (auto parameters + beta tools, DCF-style)
7) Forecast book value and earnings, compute residual income PVs, add terminal value, and output total equity value
"""
        )

    if "guide_exp_38_2__step_0___upload_statements__is___bs___soce_" not in st.session_state:
        st.session_state["guide_exp_38_2__step_0___upload_statements__is___bs___soce_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_38_2__step_0___upload_statements__is___bs___soce_"] else "▶ ") + "2) Step 0 — Upload Statements (IS + BS + SoCE)",
        key="btn_guide_exp_38_2__step_0___upload_statements__is___bs___soce_"
    ):
        st.session_state["guide_exp_38_2__step_0___upload_statements__is___bs___soce_"] = not st.session_state["guide_exp_38_2__step_0___upload_statements__is___bs___soce_"]
    if st.session_state["guide_exp_38_2__step_0___upload_statements__is___bs___soce_"]:
        st.markdown(
            r"""
### Purpose
Load your bank’s Excel statements and store them in session so values do not reset when you switch tabs.

### What you do
1) Upload an Excel file (.xlsx)
2) Select the correct sheets:
   - Income Statement sheet
   - Balance Sheet sheet
   - Statement of Changes in Equity (SoCE) sheet

### What the tool does automatically
- Cleans numeric columns (removes commas, handles brackets for negatives, strips spaces)
- Detects year columns safely (even if headers are messy or “Unnamed”)

✅ If year columns cannot be detected, the tool stops and asks you to fix the Excel headers.
"""
        )

    if "guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_" not in st.session_state:
        st.session_state["guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"] else "▶ ") + "3) Step 1 — Currency & FX Conversion (DCF-style)",
        key="btn_guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"
    ):
        st.session_state["guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"] = not st.session_state["guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"]
    if st.session_state["guide_exp_39_3__step_1___currency___fx_conversion__dcf_style_"]:
        st.markdown(
            r"""
### Purpose
Handle currency conversion consistently when statements are in **ZWG**.

### What you do
1) Choose the currency of uploaded statements:
   - **USD (already converted)** → no FX conversion is applied
   - **ZWG (convert using FX Excel)** → you must upload FX Excel

### FX Excel requirements
Your FX file must contain:
- A **Date** column
- At least one of these FX columns:
  - **Interbank**
  - **Alternative**
  - **Premium**

You then choose which FX column to use.

### Conversion rules (very important)
This model uses different FX logic for different statements:

**A) Income Statement (IS) + SoCE**
- Converted using **Yearly Average FX** for each year column.

**B) Balance Sheet (BS)**
- Converted using **Closing FX rate per year**, where you choose the closing date for each year.

This mirrors standard financial practice (flows vs stocks).

### Optional: Manual ZWG → ZiG factor (mixed periods)
If your FX history includes mixed regimes, you can enable a manual factor:
- Select year(s)
- Define date ranges within the year
- Apply a factor that divides the FX rate inside those ranges

✅ The app refreshes FX conversion automatically whenever any FX setting changes.
"""
        )

    if "guide_exp_40_4__step_2___soce_mapping__closing_equity_total_" not in st.session_state:
        st.session_state["guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"] else "▶ ") + "4) Step 2 — SoCE Mapping (Closing equity total)",
        key="btn_guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"
    ):
        st.session_state["guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"] = not st.session_state["guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"]
    if st.session_state["guide_exp_40_4__step_2___soce_mapping__closing_equity_total_"]:
        st.markdown(
            r"""
### Purpose
Tell the model where “Total Equity” is inside the SoCE.

### What you do
1) Select the **Closing Balance** row (Normalised if available)
2) Select the **TOTAL Equity column**

### Output you get
A small table showing the mapped SoCE equity totals.

⚠ Note:
SoCE mapping is mainly for reference / validation.  
The model uses **Balance Sheet equity** as the base (for consistency).
"""
        )

    if "guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_" not in st.session_state:
        st.session_state["guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"] else "▶ ") + "5) Step 3 — Balance Sheet Mapping (Equity rows)",
        key="btn_guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"
    ):
        st.session_state["guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"] = not st.session_state["guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"]
    if st.session_state["guide_exp_41_5__step_3___balance_sheet_mapping__equity_rows_"]:
        st.markdown(
            r"""
### Purpose
Define which Balance Sheet row(s) represent **Total Equity**.

### What you do
- Multi-select all rows that represent Total Equity.
  Example: if equity is split into components and there is no single “Total Equity” line,
  you can select multiple equity lines and the tool sums them.

### Why this matters
Book value (equity) is the anchor of the Residual Income model, so mapping must be correct.

✅ If you select nothing, the tool stops.
"""
        )

    if "guide_exp_42_6__step_4___earnings_line_selection__income_statement_" not in st.session_state:
        st.session_state["guide_exp_42_6__step_4___earnings_line_selection__income_statement_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_42_6__step_4___earnings_line_selection__income_statement_"] else "▶ ") + "6) Step 4 — Earnings line selection (Income Statement)",
        key="btn_guide_exp_42_6__step_4___earnings_line_selection__income_statement_"
    ):
        st.session_state["guide_exp_42_6__step_4___earnings_line_selection__income_statement_"] = not st.session_state["guide_exp_42_6__step_4___earnings_line_selection__income_statement_"]
    if st.session_state["guide_exp_42_6__step_4___earnings_line_selection__income_statement_"]:
        st.markdown(
            r"""
### Purpose
Tell the model which Income Statement line represents the “earnings” used in residual income.

### Default logic
The tool tries to default to:
- “Normalised profit / Normalized profit”

If not found, it falls back to typical lines like:
- Profit for the year
- Profit after tax (PAT)
- Net profit

### What you do
Pick the best earnings line from the dropdown.

✅ The selected earnings becomes the base earnings and forecast starting point.
"""
        )

    if "guide_exp_43_7__step_5___base_year_selection__actual_years_" not in st.session_state:
        st.session_state["guide_exp_43_7__step_5___base_year_selection__actual_years_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_43_7__step_5___base_year_selection__actual_years_"] else "▶ ") + "7) Step 5 — Base Year selection (actual years)",
        key="btn_guide_exp_43_7__step_5___base_year_selection__actual_years_"
    ):
        st.session_state["guide_exp_43_7__step_5___base_year_selection__actual_years_"] = not st.session_state["guide_exp_43_7__step_5___base_year_selection__actual_years_"]
    if st.session_state["guide_exp_43_7__step_5___base_year_selection__actual_years_"]:
        st.markdown(
            r"""
### Purpose
Select the year that will act as your “Year 0” starting point.

### How the tool chooses base-year options
- It uses the intersection of year columns available in:
  - Income Statement
  - Balance Sheet

### What you do
Choose a base year from the dropdown.

### Output you get
The tool displays:
- Total Equity (base year)
- Earnings (base year)
- Earnings line name

✅ In this model, base-year equity is always taken from the Balance Sheet.
"""
        )

    if "guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o" not in st.session_state:
        st.session_state["guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"] else "▶ ") + "8) Step 6 — Cost of Equity (Ke) via CAPM (DCF-style Auto + Override)",
        key="btn_guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"
    ):
        st.session_state["guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"] = not st.session_state["guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"]
    if st.session_state["guide_exp_44_8__step_6___cost_of_equity__ke__via_capm__dcf_style_auto___o"]:
        st.markdown(
            r"""
### Purpose
Compute **Ke** (required return on equity) using CAPM:

"""
        )
        st.latex(r"K_e = R_f + \beta \times MRP")
        st.markdown(
            r"""
### Inputs supported
**A) Country params (ERP + Default Spread)**
- Uses either:
  - default dcf_parameters.xlsx (if present), or
  - uploaded file (optional)
- Then computes:
  - **MRP = ERP**
  - **Rf = Avg Cost of Debt − Default Spread** (Zimbabwe USD assumption)

**B) Beta selection**
You can:
- Blend **unlevered betas (βu)** from selected industries (simple/weighted average), then lever it
OR
- Override directly with **manual levered beta (β)**

**Levering formula**
"""
        )
        st.latex(r"\beta_L = \beta_u \times (1 + (1 - tax)\times D/E)")
        st.markdown(
            r"""
### What you do
1) Choose whether to upload Country Params / Industry Betas (optional)
2) Select country (for ERP + spread)
3) Choose beta mode (βu then lever, or manual β)
4) Enter RF, MRP, tax and D/E (or accept auto values)

### Output you get
Rf, MRP, β, and final Ke displayed as metrics.
"""
        )

    if "guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear" not in st.session_state:
        st.session_state["guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"] else "▶ ") + "9) Step 7 — Forecast assumptions (Book Value, Discounts, Earnings growth)",
        key="btn_guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"
    ):
        st.session_state["guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"] = not st.session_state["guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"]
    if st.session_state["guide_exp_45_9__step_7___forecast_assumptions__book_value__discounts__ear"]:
        st.markdown(
            r"""
### Purpose
Define how book value and earnings evolve after the base year.

### Forecast years
You choose how many forecast years (1 to 15).  
Forecast years are: base_year+1 … base_year+n

### A) Book Value growth (YoY) + Discount
You can choose:
- **Uniform** (same rate every year)
- **Different per year**

There is also an option:
- ✅ “Auto-fill Book Value YoY (%) from BS actual YoY”
  - It computes the last actual YoY from the Balance Sheet (previous year → base year)
  - It then pre-fills the YoY input (you can still override)

### B) Earnings growth
Same structure:
- Uniform or Different per year
- Applied to the base-year earnings

### C) Terminal growth (g)
Used to compute the terminal value based on the last residual income.

✅ Tip:
Ke must be greater than terminal g, otherwise terminal value becomes invalid.
"""
        )

    if "guide_exp_46_10__model_engine___residual_income_table_and_formulas" not in st.session_state:
        st.session_state["guide_exp_46_10__model_engine___residual_income_table_and_formulas"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_46_10__model_engine___residual_income_table_and_formulas"] else "▶ ") + "10) Model engine — Residual Income table and formulas",
        key="btn_guide_exp_46_10__model_engine___residual_income_table_and_formulas"
    ):
        st.session_state["guide_exp_46_10__model_engine___residual_income_table_and_formulas"] = not st.session_state["guide_exp_46_10__model_engine___residual_income_table_and_formulas"]
    if st.session_state["guide_exp_46_10__model_engine___residual_income_table_and_formulas"]:
        st.markdown(
            r"""
### Core building blocks
For each year:

**1) Equity Charge**
"""
        )
        st.latex(r"\text{Equity Charge}_t = -K_e \times BV_t")
        st.markdown(
            r"""
**2) Residual Income**
"""
        )
        st.latex(r"RI_t = Earnings_t + \text{Equity Charge}_t")
        st.markdown(
            r"""
**3) Discount Factor**
The model supports two timing conventions:
- Base year t = 0 (standard)
- Base year t = 1 (shifted)

Standard discount factor:
"""
        )
        st.latex(r"DF_t = \frac{1}{(1+K_e)^t}")
        st.markdown(
            r"""
**4) Present Value of residual income**
"""
        )
        st.latex(r"PV(RI_t) = RI_t \times DF_t")
        st.markdown(
            r"""
### Terminal value
Terminal value is computed from the last forecast residual income:

"""
        )
        st.latex(r"TV = \frac{RI_{last}\times(1+g)}{K_e - g}")
        st.markdown(
            r"""
and present-valued using the last year discount factor.

### Final equity value
"""
        )
        st.latex(r"Equity\ Value = BV_0 + \sum PV(RI_t) + PV(TV)")
        st.markdown(
            r"""
### Output you get
1) A full “Residual Income Valuation Table (Totals)” showing:
- Book value, YoY, discounts, adjusted YoY
- Earnings and growth
- Equity charge, residual income
- Discount factors and PVs
- Terminal value and PV terminal

2) A summary with:
- Beginning Book Value
- Sum PV of residual income
- PV terminal
- Total equity value
"""
        )

    if "guide_exp_47_11__outputs_saved_for_other_tabs__integration_" not in st.session_state:
        st.session_state["guide_exp_47_11__outputs_saved_for_other_tabs__integration_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_47_11__outputs_saved_for_other_tabs__integration_"] else "▶ ") + "11) Outputs saved for other tabs (integration)",
        key="btn_guide_exp_47_11__outputs_saved_for_other_tabs__integration_"
    ):
        st.session_state["guide_exp_47_11__outputs_saved_for_other_tabs__integration_"] = not st.session_state["guide_exp_47_11__outputs_saved_for_other_tabs__integration_"]
    if st.session_state["guide_exp_47_11__outputs_saved_for_other_tabs__integration_"]:
        st.markdown(
            r"""
This module stores key outputs into session_state so other tabs can use them:

- bank.outputs.book_equity_0  → Beginning Book Value
- bank.outputs.earnings_0     → Base earnings
- bank.outputs.ke            → Cost of equity
- bank.outputs.equity_value_total → Final equity value
- equity_value_banking       → Final equity value (shortcut)

If you return to Comparables, P/B can auto-use Book Equity from here.
"""
        )

    if "guide_exp_48_12__troubleshooting__common_issues_" not in st.session_state:
        st.session_state["guide_exp_48_12__troubleshooting__common_issues_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_48_12__troubleshooting__common_issues_"] else "▶ ") + "12) Troubleshooting (common issues)",
        key="btn_guide_exp_48_12__troubleshooting__common_issues_"
    ):
        st.session_state["guide_exp_48_12__troubleshooting__common_issues_"] = not st.session_state["guide_exp_48_12__troubleshooting__common_issues_"]
    if st.session_state["guide_exp_48_12__troubleshooting__common_issues_"]:
        st.markdown(
            r"""
### “Could not detect year columns”
- Your Excel headers may be merged or labelled “Unnamed”.
- Ensure each year appears clearly in the column header (e.g., 2022, 2023, 2024).

### FX conversion stops / missing dates
- Make sure FX file has:
  - Date column, and selected FX column is numeric
- If a Balance Sheet closing date is before the first FX date → no “as of” rate exists.

### Equity mapping gives wrong totals
- Confirm you selected the correct BS equity lines.
- If equity is split across multiple lines, select all relevant lines.

### Terminal value becomes blank/invalid
- Check that Ke > terminal g.
- If not, lower g or revise Ke assumptions.
"""
        )
# -------------------------
# SUMMARY TAB (TAB5)
# -------------------------
with tab5:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">SUMMARY</span> Summary Valuation — Weighted Equity Value</h3>
          <p class="subtle">
            This page combines outputs from <b>DCF</b>, <b>DDM</b>, <b>Comparables</b> (EV/EBITDA · PBV · P/E),
            and <b>Banking</b> into one <b>blended equity value</b>. You select models, assign weights, and the app
            calculates intrinsic value per share, upside/downside, and a simple recommendation.
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if "guide_exp_49_1__what_the_summary_page_does__big_picture_" not in st.session_state:
        st.session_state["guide_exp_49_1__what_the_summary_page_does__big_picture_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_49_1__what_the_summary_page_does__big_picture_"] else "▶ ") + "1) What the Summary page does (big picture)",
        key="btn_guide_exp_49_1__what_the_summary_page_does__big_picture_"
    ):
        st.session_state["guide_exp_49_1__what_the_summary_page_does__big_picture_"] = not st.session_state["guide_exp_49_1__what_the_summary_page_does__big_picture_"]
    if st.session_state["guide_exp_49_1__what_the_summary_page_does__big_picture_"]:
        st.markdown(
            r"""
### Purpose
The Summary page is your “final dashboard”. It:
- Pulls **equity values** already calculated in other tabs
- Lets you **choose which models to include**
- Lets you **assign weights (%)** to each selected model
- Produces a **Weighted Equity Value (blended valuation)**
- Converts the blended equity into:
  - **Intrinsic value per share** (if shares are provided)
  - **Upside/Downside %** vs current market price (if price is provided)
  - A simple **Buy / Hold / Reduce** label

### Models supported
- **DCF**
- **DDM**
- **EV/EBITDA**
- **PBV**
- **P/E**
- **BANKING**
"""
        )

    if "guide_exp_50_2__step_1___select_models_to_include" not in st.session_state:
        st.session_state["guide_exp_50_2__step_1___select_models_to_include"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_50_2__step_1___select_models_to_include"] else "▶ ") + "2) Step 1 — Select Models to Include",
        key="btn_guide_exp_50_2__step_1___select_models_to_include"
    ):
        st.session_state["guide_exp_50_2__step_1___select_models_to_include"] = not st.session_state["guide_exp_50_2__step_1___select_models_to_include"]
    if st.session_state["guide_exp_50_2__step_1___select_models_to_include"]:
        st.markdown(
            r"""
### What you do
Use the multi-select box to choose which models are “active”.

### What the app does
- Stores your selection in session_state so it won’t reset when you switch tabs.
- If you select nothing, the page stops (because we need at least one model to compute a blend).

✅ Tip:
If a model value is missing (because you didn’t run that model tab yet), it may show as blank/None in the table.
So always run the model tabs first if you want them included.
"""
        )

    if "guide_exp_51_3__step_2___assign_weights____" not in st.session_state:
        st.session_state["guide_exp_51_3__step_2___assign_weights____"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_51_3__step_2___assign_weights____"] else "▶ ") + "3) Step 2 — Assign Weights (%)",
        key="btn_guide_exp_51_3__step_2___assign_weights____"
    ):
        st.session_state["guide_exp_51_3__step_2___assign_weights____"] = not st.session_state["guide_exp_51_3__step_2___assign_weights____"]
    if st.session_state["guide_exp_51_3__step_2___assign_weights____"]:
        st.markdown(
            r"""
### What you do
For each model you selected, enter a weight (%).
- You can enter any numbers (they do not need to sum to 100).

### What the app does (important)
It automatically **normalizes** weights so the selected models sum to **100%**.

Example:
- DCF = 40, DDM = 20, PBV = 20  → total input = 80  
Normalized:
- DCF = 40/80 = 50%
- DDM = 20/80 = 25%
- PBV = 20/80 = 25%

✅ If total weight for selected models is 0, the page stops (division by zero).
"""
        )

    if "guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping" not in st.session_state:
        st.session_state["guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"] else "▶ ") + "4) Where the Summary values come from (session_state mapping)",
        key="btn_guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"
    ):
        st.session_state["guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"] = not st.session_state["guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"]
    if st.session_state["guide_exp_52_4__where_the_summary_values_come_from__session_state_mapping"]:
        st.markdown(
            r"""
### The Summary page pulls results from other tabs using these keys:

- **DCF** → `equity_value_dcf`
- **DDM** → `equity_value_ddm`
- **EV/EBITDA** → `value_ev_ebitda`
- **PBV** → `value_pbv`
- **P/E** → `value_pe`
- **BANKING** → `equity_value_banking`

If a key is missing (because you haven’t run that model yet), the value may show as blank.
"""
        )

    if "guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated" not in st.session_state:
        st.session_state["guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"] else "▶ ") + "5) How the blended (Weighted) Equity Value is calculated",
        key="btn_guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"
    ):
        st.session_state["guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"] = not st.session_state["guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"]
    if st.session_state["guide_exp_53_5__how_the_blended__weighted__equity_value_is_calculated"]:
        st.markdown(
            r"""
### For each selected model
The app computes:

**Weighted Value = Model Value × (Normalized Weight / 100)**

### Final blended equity
The **Weighted Equity Value** is:

**Weighted Equity = SUM of all Weighted Values**

This is displayed as a KPI card at the top and shown again in the table.
"""
        )

    if "guide_exp_54_6__summary_table___interactive_dashboard" not in st.session_state:
        st.session_state["guide_exp_54_6__summary_table___interactive_dashboard"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_54_6__summary_table___interactive_dashboard"] else "▶ ") + "6) Summary Table & Interactive Dashboard",
        key="btn_guide_exp_54_6__summary_table___interactive_dashboard"
    ):
        st.session_state["guide_exp_54_6__summary_table___interactive_dashboard"] = not st.session_state["guide_exp_54_6__summary_table___interactive_dashboard"]
    if st.session_state["guide_exp_54_6__summary_table___interactive_dashboard"]:
        st.markdown(
            r"""
### Summary Table tab
Shows a table with:
- Model
- Value (USD)
- Weight (%) (normalized)
- Weighted Value (USD)

### Interactive Dashboard tab
Shows bar charts for:
- Model equity values
- Model weights
- Weighted contributions (how much each model contributes to the final blended equity)

✅ Use the dashboard to quickly spot:
- which model is driving the valuation most,
- which model is the outlier (highest/lowest),
- and how wide the valuation range is.
"""
        )

    if "guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs" not in st.session_state:
        st.session_state["guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"] else "▶ ") + "7) Valuation Summary (Shares, Price, Intrinsic, Upside/Downside)",
        key="btn_guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"
    ):
        st.session_state["guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"] = not st.session_state["guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"]
    if st.session_state["guide_exp_55_7__valuation_summary__shares__price__intrinsic__upside_downs"]:
        st.markdown(
            r"""
### What you do
Enter:
1) **Number of Shares in Issue**
2) **Current Share Price (USD)**

### What the app computes
**Intrinsic Value per Share**
- Intrinsic = Weighted Equity / Shares  
(only computed if shares > 0)

**Upside/Downside (%)**
- Upside% = (Intrinsic − Current Price) / Current Price × 100  
(only computed if current price > 0)

### Output you get
A small table containing:
- Weighted Equity Value
- Shares
- Intrinsic value per share
- Current share price
- Upside / Downside (%)
"""
        )

    if "guide_exp_56_8__recommendation_logic__buy___hold___reduce_" not in st.session_state:
        st.session_state["guide_exp_56_8__recommendation_logic__buy___hold___reduce_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_56_8__recommendation_logic__buy___hold___reduce_"] else "▶ ") + "8) Recommendation logic (Buy / Hold / Reduce)",
        key="btn_guide_exp_56_8__recommendation_logic__buy___hold___reduce_"
    ):
        st.session_state["guide_exp_56_8__recommendation_logic__buy___hold___reduce_"] = not st.session_state["guide_exp_56_8__recommendation_logic__buy___hold___reduce_"]
    if st.session_state["guide_exp_56_8__recommendation_logic__buy___hold___reduce_"]:
        st.markdown(
            r"""
### The recommendation is based on Upside/Downside (%)

If shares and price are entered correctly, the app labels:

- **🟢 BUY / ACCUMULATE** if upside is meaningfully positive (above the buy threshold)
- **🟡 HOLD / FAIRLY VALUED** if upside is near zero (within a fair value band)
- **🔴 REDUCE / SELL** if upside is negative beyond the band

✅ Note:
This recommendation is purely rule-based and depends heavily on:
- model weights,
- your assumptions in each valuation tab,
- Ke/WACC/growth choices,
- and data quality.
"""
        )

    if "guide_exp_57_9__excel_download___valuation_summary__with_formulas_" not in st.session_state:
        st.session_state["guide_exp_57_9__excel_download___valuation_summary__with_formulas_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_57_9__excel_download___valuation_summary__with_formulas_"] else "▶ ") + "9) Excel Download — Valuation Summary (with formulas)",
        key="btn_guide_exp_57_9__excel_download___valuation_summary__with_formulas_"
    ):
        st.session_state["guide_exp_57_9__excel_download___valuation_summary__with_formulas_"] = not st.session_state["guide_exp_57_9__excel_download___valuation_summary__with_formulas_"]
    if st.session_state["guide_exp_57_9__excel_download___valuation_summary__with_formulas_"]:
        st.markdown(
            r"""
### What you get in the Excel export
When you click download, the file includes:

**Sheet 1: Model_Summary**
- Model
- Value_USD
- Weight_Input_%
- Weight_Normalized_% (formula)
- Weighted_Value_USD (formula)
- Total blended equity (SUM formula)

**Sheet 2: Valuation_Summary**
- Weighted Equity Value (linked from Sheet 1)
- Shares (your input)
- Intrinsic value per share (formula)
- Current price (your input)
- Upside/Downside % (formula)
- Recommendation (formula)

✅ This makes the export easy to audit and share with stakeholders.
"""
        )

    if "guide_exp_58_10__troubleshooting__common_issues_" not in st.session_state:
        st.session_state["guide_exp_58_10__troubleshooting__common_issues_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_58_10__troubleshooting__common_issues_"] else "▶ ") + "10) Troubleshooting (common issues)",
        key="btn_guide_exp_58_10__troubleshooting__common_issues_"
    ):
        st.session_state["guide_exp_58_10__troubleshooting__common_issues_"] = not st.session_state["guide_exp_58_10__troubleshooting__common_issues_"]
    if st.session_state["guide_exp_58_10__troubleshooting__common_issues_"]:
        st.markdown(
            r"""
### “My model value is blank / —”
- You probably have not run that model tab yet.
- Go to the model tab (DCF / DDM / Comparables / Banking), complete it, then return.

### “Total weight cannot be zero”
- You selected models, but typed 0 for all their weights.
- Give at least one selected model a non-zero weight.

### Intrinsic value shows blank/NaN
- Shares must be > 0.

### Upside/Downside shows blank/NaN
- Current price must be > 0 (and shares must be valid so intrinsic exists).

### Excel export downloads but values look wrong
- Check that your selected models have real values.
- Check you did not accidentally set weights to 0 or exclude a key model.
"""
        )
# -------------------------
# TAB6 🛠 TROUBLESHOOTING (ALL MODELS)
# -------------------------
with tab6:
    st.markdown(
        """
        <div class="card">
          <h3><span class="pill">HELP</span> 🛠 Troubleshooting — All Models</h3>
          <p class="subtle">
            Use this section when something looks wrong, missing, blank, or “not updating”.
            Most issues are caused by missing inputs, missing source files, wrong sheet selection,
            or session state resets after clearing.
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -------------------------
    # QUICK DIAGNOSTIC CHECKLIST
    # -------------------------
    if "guide_exp_59_0__quick_checklist__do_this_first_" not in st.session_state:
        st.session_state["guide_exp_59_0__quick_checklist__do_this_first_"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_59_0__quick_checklist__do_this_first_"] else "▶ ") + "0) Quick checklist (do this first)",
        key="btn_guide_exp_59_0__quick_checklist__do_this_first_"
    ):
        st.session_state["guide_exp_59_0__quick_checklist__do_this_first_"] = not st.session_state["guide_exp_59_0__quick_checklist__do_this_first_"]
    if st.session_state["guide_exp_59_0__quick_checklist__do_this_first_"]:
        st.markdown(
            r"""
### Quick checklist (90% of issues)
1) **Check you ran the model tab first**  
   - Summary pulls outputs from each model’s `session_state` keys.  
   - If you never ran DCF/DDM/Comparables/Banking, Summary will show blanks.

2) **Look for red/amber warnings on the page**
   - Those warnings tell you exactly which input or file is missing.

3) **Avoid clearing a tab accidentally**
   - “Clear” buttons wipe that tab’s session values.

4) **Check currency & FX settings**
   - Wrong FX conversion can make values look 10×, 100×, or 1000× too big.

5) **Check year detection**
   - If your Excel has merged headers (“Unnamed: …”), the app may fail to detect years.
"""
        )

    # -------------------------
    # DCF TROUBLESHOOTING
    # -------------------------
    if "guide_exp_60_2__dcf___common_issues" not in st.session_state:
        st.session_state["guide_exp_60_2__dcf___common_issues"] = False
    if st.button(
        ("▼ " if st.session_state["guide_exp_60_2__dcf___common_issues"] else "▶ ") + "2) DCF — common issues",
        key="btn_guide_exp_60_2__dcf___common_issues"
    ):
        st.session_state["guide_exp_60_2__dcf___common_issues"] = not st.session_state["guide_exp_60_2__dcf___common_issues"]
    if st.session_state["guide_exp_60_2__dcf___common_issues"]:
        st.markdown(
            r"""
### Problem: “DCF says missing data / UFCF is empty”
**Common causes**
- You didn’t load the financial inputs correctly.
- You didn’t complete required assumptions (WACC, growth, tax, margins, etc.).
- Forecast years are not aligned.

**Fix**
- Confirm the DCF input sheet has required line items and valid year columns.
- Ensure WACC and terminal growth assumptions are valid.
- Check any warnings about missing revenue/EBIT/FCF drivers.

### Problem: “Terminal value is NaN / error”
**Cause**
- Terminal growth `g` is ≥ WACC (or Ke), making denominator invalid.
**Fix**
- Reduce `g` or increase discount rate so that **WACC > g**.
"""
        )

    # -------------------------
    # DDM TROUBLESHOOTING
    # -------------------------
    if "guide_exp_61_3__ddm___common_issues" not in st.session_state:
        st.session_state["guide_exp_61_3__ddm___common_issues"] = False
    if st.button(
        ("▼ " if st.session_state["guide_exp_61_3__ddm___common_issues"] else "▶ ") + "3) DDM — common issues",
        key="btn_guide_exp_61_3__ddm___common_issues"
    ):
        st.session_state["guide_exp_61_3__ddm___common_issues"] = not st.session_state["guide_exp_61_3__ddm___common_issues"]
    if st.session_state["guide_exp_61_3__ddm___common_issues"]:
        st.markdown(
            r"""
### Problem: “DDM Equity is NaN / not computed”
**Common causes**
- Dividends are missing or 0.
- `Ke` is missing or <= terminal dividend growth.

**Fix**
- Ensure dividends (D1 / forecast dividends) exist and are numeric.
- Ensure **Ke > g** for terminal dividend growth.

### Problem: “DDM value is extremely small”
**Cause**
- Dividend inputs are tiny relative to shares/equity.
**Fix**
- Validate dividend units and currency (USD vs ZWG conversion).
"""
        )

    # -------------------------
    # COMPARABLES TROUBLESHOOTING
    # -------------------------
    if "guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_" not in st.session_state:
        st.session_state["guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"] = False
    if st.button(
        ("▼ " if st.session_state["guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"] else "▶ ") + "4) Comparables — common issues (EV/EBITDA, PBV, P/E)",
        key="btn_guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"
    ):
        st.session_state["guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"] = not st.session_state["guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"]
    if st.session_state["guide_exp_62_4__comparables___common_issues__ev_ebitda__pbv__p_e_"]:
        st.markdown(
            r"""
### Problem: “Suggested peers look wrong”
**Cause**
- PeerGroup / Sector / Industry mapping is wrong in `peer_universe.xlsx`.
**Fix**
- Check your target company row has the right PeerGroup.
- If PeerGroup is present, the app uses STRICT PeerGroup matching.

### Problem: “Averages look wrong”
**Cause**
- You included/excluded the wrong peers using the Analyst Filter (Include EV/PB/PE).
- You left 0s in multiples (0s are treated as real values, not blanks).

**Fix**
- Use the Include checkboxes carefully (turn off outliers or irrelevant comps).
- Replace “unknown” values with blanks (NaN) rather than 0.

### Problem: “Maintainable EBITDA/Earnings is missing”
**Cause**
- Comparables Step 3 & Step 4 depend on DCF outputs:
  - EBITDA: `dcf_ebitda_all` / `dcf_ebitda_forecast`
  - Earnings: `dcf_profit_all` / `dcf_profit_forecast`

**Fix**
- Run DCF first, then come back to Comparables.
- If DCF has no valid year keys (must be 4-digit years like 2024), fix the DCF output structure.

### Problem: “Excel export formulas don’t work”
**Cause**
- Excel formulas rely on ranges (start/end rows) and include flags.
**Fix**
- Ensure you have at least 1 comparable row in Comps_Input.
- Check Include flags are TRUE/FALSE in Excel (not text).
"""
        )

    # -------------------------
    # BANKING TROUBLESHOOTING
    # -------------------------
    if "guide_exp_63_5__banking___common_issues__residual_income_" not in st.session_state:
        st.session_state["guide_exp_63_5__banking___common_issues__residual_income_"] = False
    if st.button(
        ("▼ " if st.session_state["guide_exp_63_5__banking___common_issues__residual_income_"] else "▶ ") + "5) Banking — common issues (Residual Income)",
        key="btn_guide_exp_63_5__banking___common_issues__residual_income_"
    ):
        st.session_state["guide_exp_63_5__banking___common_issues__residual_income_"] = not st.session_state["guide_exp_63_5__banking___common_issues__residual_income_"]
    if st.session_state["guide_exp_63_5__banking___common_issues__residual_income_"]:
        st.markdown(
            r"""
### Problem: “Year columns not detected”
**Cause**
- Your Excel has merged headers or “Unnamed: …” columns.
**Fix**
- Make sure year headers contain a real year (e.g., 2022, 2023) in the column label.
- Avoid merged year cells in the header row.

### Problem: “FX conversion stopped / missing Date column”
**Cause**
- FX Excel must have **Date** plus one of: **Interbank / Alternative / Premium**.
**Fix**
- Add a `Date` column (Excel date format) and at least one allowed FX column.

### Problem: “Equity / Earnings is wrong”
**Cause**
- Wrong sheet selection (IS vs BS vs SoCE).
- Wrong row mapping for Equity or Earnings.
**Fix**
- Re-check sheet selections at upload step.
- Re-select correct Equity rows (multi-select) and Earnings line.

### Problem: “Terminal value is NaN”
**Cause**
- `Ke <= g_term`.
**Fix**
- Ensure **Ke > terminal growth**.
"""
        )

    # -------------------------
    # SUMMARY TROUBLESHOOTING
    # -------------------------
    if "guide_exp_64_6__summary___common_issues__weighted_equity_" not in st.session_state:
        st.session_state["guide_exp_64_6__summary___common_issues__weighted_equity_"] = False
    if st.button(
        ("▼ " if st.session_state["guide_exp_64_6__summary___common_issues__weighted_equity_"] else "▶ ") + "6) Summary — common issues (weighted equity)",
        key="btn_guide_exp_64_6__summary___common_issues__weighted_equity_"
    ):
        st.session_state["guide_exp_64_6__summary___common_issues__weighted_equity_"] = not st.session_state["guide_exp_64_6__summary___common_issues__weighted_equity_"]
    if st.session_state["guide_exp_64_6__summary___common_issues__weighted_equity_"]:
        st.markdown(
            r"""
### Problem: “A model shows blank/None”
**Cause**
- You did not run that model tab yet (so session_state key is missing).
**Fix**
- Run the model tab first, then return to Summary.

### Problem: “Total weight cannot be zero”
**Fix**
- Give at least one selected model a non-zero weight.

### Problem: “Intrinsic value is blank”
**Fix**
- Shares must be > 0.

### Problem: “Upside/Downside is blank”
**Fix**
- Current price must be > 0 (and shares must be valid).
"""
        )

    # -------------------------
    # DATA QUALITY & UNITS
    # -------------------------
    if "guide_exp_65_7__data_quality__units___currency_problems__big_values___tin" not in st.session_state:
        st.session_state["guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"] = True
    if st.button(
        ("▼ " if st.session_state["guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"] else "▶ ") + "7) Data quality, units & currency problems (big values / tiny values)",
        key="btn_guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"
    ):
        st.session_state["guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"] = not st.session_state["guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"]
    if st.session_state["guide_exp_65_7__data_quality__units___currency_problems__big_values___tin"]:
        st.markdown(
            r"""
### Problem: “Values are 10× / 100× / 1000× too big/small”
This is almost always **units or FX**.

**Check:**
- Are statements in USD already?  
  If yes, do **NOT** apply FX conversion.
- If converting ZWG → USD:
  - Make sure the FX rate column is correct
  - Check if you applied the ZiG factor correctly (only in relevant date ranges)
- Confirm whether your statements are in:
  - dollars, thousands, or millions
  If statements are “in thousands”, your valuation must be scaled consistently.

✅ Practical check:
Compare one known line item (e.g., Total Equity) to an annual report figure to confirm scale.
"""
        )
# -------------------------
# TAB7 ⚡ QUICK SUMMARY (FAST HOW-TO)
# -------------------------
with tab7:
    if "guide_exp_0_1__what_the_dcf_page_does__big_picture_" not in st.session_state:
        st.session_state["guide_exp_0_1__what_the_dcf_page_does__big_picture_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_0_1__what_the_dcf_page_does__big_picture_"] else "▶ ") + "1) What the DCF page does (big picture)",
            key="btn_guide_exp_0_1__what_the_dcf_page_does__big_picture_"
    ):
        st.session_state["guide_exp_0_1__what_the_dcf_page_does__big_picture_"] = not st.session_state[
            "guide_exp_0_1__what_the_dcf_page_does__big_picture_"]
    if st.session_state["guide_exp_0_1__what_the_dcf_page_does__big_picture_"]:
        st.markdown("""
       **This DCF page produces an Equity Value using an Unlevered Free Cash Flow (UFCF/FCFF) model.**

       It takes:
       - **Income Statement + Balance Sheet + Cash Flow** (historical),
       - converts currency to USD if needed,
       - lets you **map the correct lines** (Debt, Cash, Revenue, EBITDA, etc.),
       - forecasts the Income Statement and Working Capital,
       - computes **WACC** using CAPM + cost of debt,
       - discounts UFCF + Terminal Value,
       - outputs **Enterprise Value → Equity Value**,
       - and builds a **WACC vs g sensitivity grid** + a downloadable Excel model.

       **Final formula logic**
       - **EV = PV(UFCF forecast years) + PV(Terminal Value)**
       - **Equity Value = EV − Net Debt**
       """)

    if "guide_exp_1_2__step_0___start_new_valuation__reset_button_" not in st.session_state:
        st.session_state["guide_exp_1_2__step_0___start_new_valuation__reset_button_"] = False
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_1_2__step_0___start_new_valuation__reset_button_"] else "▶ ") + "2) Step 0 — Start New Valuation (Reset button)",
            key="btn_guide_exp_1_2__step_0___start_new_valuation__reset_button_"
    ):
        st.session_state["guide_exp_1_2__step_0___start_new_valuation__reset_button_"] = not st.session_state[
            "guide_exp_1_2__step_0___start_new_valuation__reset_button_"]
    if st.session_state["guide_exp_1_2__step_0___start_new_valuation__reset_button_"]:
        st.markdown("""
       At the top of the DCF page you will see:

       ✅ **🗂️ Clear & Upload New File**

       Use this when you want to start a fresh valuation.

       It clears:
       - uploaded statements & FX files
       - all mappings (Revenue, Debt, Cash, etc.)
       - forecasts, WACC inputs, sensitivity settings, and outputs

       **Tip:** If you uploaded the wrong Excel or mapped wrong rows, click reset first.
       """)

    if "guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req" not in st.session_state:
        st.session_state["guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"] else "▶ ") + "3) Step 1 — Upload Financial Statements (Excel structure required)",
            key="btn_guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"
    ):
        st.session_state["guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"] = not \
        st.session_state["guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"]
    if st.session_state["guide_exp_2_3__step_1___upload_financial_statements__excel_structure_req"]:
        st.markdown("""
       ### ✅ Required file format
       Upload **ONE Excel file** with **3 sheets in this exact order**:

       1. **Sheet 1:** Income Statement (IS)  
       2. **Sheet 2:** Balance Sheet (BS)  
       3. **Sheet 3:** Cash Flow (CF)

       ### ✅ Required layout inside each sheet
       - The **first column must be line items** (e.g., Revenue, EBITDA, Total Assets)
       - All other columns must be **years** (e.g., 2022, 2023, 2024)
       - Numbers can include commas and brackets (e.g. `(1,250)`), the system cleans them.

       ### Common upload mistakes
       - If your years are typed like `FY2024` instead of `2024`, results may fail.
       - If the first column is not line items, mapping will be confusing.
       """)

    if "guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_" not in st.session_state:
        st.session_state["guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"] else "▶ ") + "4) Step 2 — Currency & FX Conversion (USD vs ZWG)",
            key="btn_guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"
    ):
        st.session_state["guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"] = not st.session_state[
            "guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"]
    if st.session_state["guide_exp_3_4__step_2___currency___fx_conversion__usd_vs_zwg_"]:
        st.markdown("""
       ### 4.1 Choose the currency of uploaded statements
       You must choose one:
       - ✅ **USD (already converted)** → no FX conversion is done
       - ⚠️ **ZWG (convert using FX Excel)** → you must upload FX data

       ### 4.2 If ZWG: Upload FX Excel
       Your FX file must have:
       - A **Date** column
       - At least one of these rate columns:
         - **Interbank**
         - **Alternative**
         - **Premium**

       You will select which FX column to use.

       ### 4.3 How FX conversion is applied
       The model converts ZWG → USD like this:

       - **Income Statement & Cash Flow**: uses **Yearly Average FX**  
       - **Balance Sheet**: uses **Closing FX per year** based on the date you choose

       That is correct finance logic because:
       - IS/CF flows happen throughout the year → average rate is reasonable  
       - BS is a point-in-time snapshot → use closing rate

       ### 4.4 Balance Sheet closing dates (per year)
       You will enter the closing date for each Balance Sheet year (default is 31 Dec).
       The system will pick the last FX rate available **on or before** that date.

       ### 4.5 Optional: ZWG → ZiG factor adjustment
       If your data spans a period where ZWG was replaced or a factor is needed:
       - enable the factor,
       - select the year(s),
       - select date ranges,
       - the system divides the FX values by your factor inside those ranges.

       **Use this only if you truly have mixed periods.**
       """)

    if "guide_exp_4_5__step_3___mapping__most_important_step_" not in st.session_state:
        st.session_state["guide_exp_4_5__step_3___mapping__most_important_step_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_4_5__step_3___mapping__most_important_step_"] else "▶ ") + "5) Step 3 — Mapping (MOST IMPORTANT STEP)",
            key="btn_guide_exp_4_5__step_3___mapping__most_important_step_"
    ):
        st.session_state["guide_exp_4_5__step_3___mapping__most_important_step_"] = not st.session_state[
            "guide_exp_4_5__step_3___mapping__most_important_step_"]
    if st.session_state["guide_exp_4_5__step_3___mapping__most_important_step_"]:
        st.markdown("""
       Mapping means: **you tell the model which statement lines represent the variables it needs.**  
       If mapping is wrong, the valuation will be wrong.

       You will map in 3 places:

       ✅ **A) Balance Sheet Mapping**  
       ✅ **B) Cash Flow Mapping**  
       ✅ **C) Income Statement Core Totals Mapping**
       """)

        st.markdown("### 5A) Balance Sheet Mapping — what to select")
        st.markdown("""
       You will multi-select rows (you can select more than one if statements are split).

       **1) Total Debt / Borrowings**  
       Select all interest-bearing debt rows, e.g.:
       - Loans and borrowings  
       - Bank loans  
       - Notes / bonds  
       - Lease liabilities (if you want)  
       Avoid: trade payables (those are working capital).

       **2) Cash & Cash Equivalents**  
       Select:
       - Cash, bank balances, cash equivalents, short-term deposits.

       **3) Current Assets (CA) for Working Capital**  
       Select operating current assets like:
       - Inventory  
       - Trade receivables / Debtors  
       - Prepayments (optional)

       Avoid: cash if you already mapped cash separately (unless your company treats cash as operating).

       **4) Current Liabilities (CL) for Working Capital**  
       Select operating current liabilities like:
       - Trade payables / Creditors  
       - Accrued expenses  
       - Other payables

       Avoid: interest-bearing short-term debt if you already include it in “Debt”.

       **5) Equity (Book Equity)**
       This is the equity used for **D/E (Debt-to-Equity)** in your WACC calculation.
       Select:
       - Total equity  
       - Shareholders’ equity  
       - Equity attributable to owners  
       If there are multiple equity lines, you can multi-select them.

       ✅ **Important:** This is **book equity from the balance sheet**, not market cap.
       """)

        st.markdown("### 5B) Cash Flow Mapping — what to select")
        st.markdown("""
       **1) Depreciation & Amortisation (CF)**  
       Select depreciation line in Cash Flow if it exists.

       **2) Capex (CF)**  
       Select capex type lines like:
       - Purchase of property, plant and equipment (PPE)
       - Additions to PPE
       - Purchase of intangibles

       Capex is often negative (cash outflow). The model keeps sign.

       **3) Interest paid (optional)**
       Only select if your Income Statement doesn’t clearly contain interest expense.
       This helps compute **cost of debt**.
       """)

        st.markdown("### 5C) Income Statement Core Totals Mapping — what to select")
        st.markdown("""
       This is a step-by-step wizard. Only **Revenue is mandatory**, but better mapping improves accuracy.

       **Revenue (MANDATORY)**  
       Select total revenue / sales.

       Optional but recommended:
       - Cost of Sales / Raw Materials
       - Gross Profit
       - EBITDA
       - Depreciation & Amortisation (IS line)  ✅ (your code supports this)
       - Operating Profit / EBIT
       - Profit Before Tax (PBT)
       - Income Tax (tax expense)
       - Profit for the Year (Net profit)

       ✅ The model checks that totals appear **top-to-bottom** in correct order.
       If you map totals out of order it will stop and show an error.
       """)

    if "guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth" not in st.session_state:
        st.session_state["guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"] else "▶ ") + "6) Step 4 — Forecast horizon and revenue growth",
            key="btn_guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"
    ):
        st.session_state["guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"] = not st.session_state[
            "guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"]
    if st.session_state["guide_exp_5_6__step_4___forecast_horizon_and_revenue_growth"]:
        st.markdown("""
       ### 6.1 Forecast horizon
       Choose how many years to forecast (1 to 15).

       ### 6.2 Revenue growth
       The system calculates a historical average revenue growth and shows it.
       You can override it.

       ### 6.3 Growth method choice
       You can choose:
       - **Uniform growth** (same % every forecast year)
       - **Different growth per year** (enter each year separately)

       Use “Different growth” if you want realistic fade down/up patterns.
       """)

    if "guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall" not in st.session_state:
        st.session_state["guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"] else "▶ ") + "7) Step 5 — Forecast logic (what the model does automatically)",
            key="btn_guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"
    ):
        st.session_state["guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"] = not \
        st.session_state["guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"]
    if st.session_state["guide_exp_6_7__step_5___forecast_logic__what_the_model_does_automaticall"]:
        st.markdown("""
       After revenue is forecasted, the system forecasts other IS lines.

       ### 7.1 Gross Profit / COS handling
       There are 4 cases:
       - **GP and COS mapped:** forecasts COS using average GP margin  
       - **GP mapped, COS missing:** forecasts GP using average GP margin  
       - **COS mapped, GP missing:** forecasts COS as % of revenue  
       - **Neither mapped:** no special handling; other rows still forecast as % of revenue

       ### 7.2 “Other rows as % of revenue”
       For every non-total line item not protected:
       - It computes an average historical ratio (Row / Revenue)
       - Forecasts the row using that ratio × forecast revenue

       ### 7.3 Totals chain engine
       Totals like GP, EBITDA, EBIT, PBT, NP are re-calculated by summing the block between totals.
       This prevents double-counting and keeps totals consistent.

       ### 7.4 Tax forecasting
       Tax is derived using:
       - average historical **Tax / PBT** ratio (only when PBT is positive)
       Tax stays negative if your statement shows tax as negative.
       """)

    if "guide_exp_7_8__step_6___working_capital__historical___wc____forecast____" not in st.session_state:
        st.session_state["guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"] else "▶ ") + "8) Step 6 — Working Capital (Historical → WC% → Forecast → ΔWC)",
            key="btn_guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"
    ):
        st.session_state["guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"] = not \
        st.session_state["guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"]
    if st.session_state["guide_exp_7_8__step_6___working_capital__historical___wc____forecast____"]:
        st.markdown("""
       This section only works if you mapped **Current Assets and Current Liabilities**.

       ### 8.1 Historical Working Capital
       - WC = CA − CL (by year)

       ### 8.2 WC % of Sales
       - WC% = WC / Revenue

       You can **exclude outlier years** using the “Include” checkbox table.
       This is important if one year is abnormal.

       ### 8.3 Choose assumption method
       You choose:
       - **Average WC%** (across included years)
       OR
       - **Most recent WC%** (latest included year)

       ### 8.4 Forecast WC
       - Forecast WC = Forecast Revenue × WC%

       ### 8.5 Change in WC (ΔWC)
       Your model defines:
       - **ΔWC = Old WC − New WC**

       Meaning:
       - If WC increases → ΔWC becomes negative (cash outflow)
       - If WC decreases → ΔWC becomes positive (cash inflow)
       """)

    if "guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio" not in st.session_state:
        st.session_state["guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"] else "▶ ") + "9) Step 7 — Debt, Cash, Net Debt and D/E ratio",
            key="btn_guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"
    ):
        st.session_state["guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"] = not st.session_state[
            "guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"]
    if st.session_state["guide_exp_8_9__step_7___debt__cash__net_debt_and_d_e_ratio"]:
        st.markdown("""
       These come from the **Balance Sheet mapping** at the last common year.

       - **Total Debt** = sum of selected debt rows
       - **Cash** = sum of selected cash rows
       - **Net Debt** = Debt − Cash
       - **Equity** = sum of selected equity rows
       - **D/E ratio** = Debt / Equity

       ✅ D/E is used to compute weights in WACC:
       - wd = D/E / (1 + D/E)
       - we = 1 / (1 + D/E)

       ⚠️ If your equity mapping is wrong, your D/E and WACC will be wrong.
       """)

    if "guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_" not in st.session_state:
        st.session_state["guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"] else "▶ ") + "10) Step 8 — DCF Parameters (RF, MRP, Beta, Tax, Rd, g)",
            key="btn_guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"
    ):
        st.session_state["guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"] = not st.session_state[
            "guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"]
    if st.session_state["guide_exp_9_10__step_8___dcf_parameters__rf__mrp__beta__tax__rd__g_"]:
        st.markdown("""
       ### 10.1 Risk-free rate (RF) + Market risk premium (MRP)
       You can run in two ways:
       - **Auto (from Excel)** using Country ERP + Default Spread file
       - **Manual override** (type in values)

       If auto is ON and the Excel has values, the model will “snap” RF & MRP to the auto results.

       ### 10.2 Unlevered beta (βu)
       You can:
       - Select one or more industries from the Industry Betas file
       - Choose:
         - Simple average
         - Weighted average
       - Or manually override βu

       ### 10.3 Cost of debt (Rd)
       Auto mode:
       - Rd = |Interest| / |Total Debt|  (from statements)

       Manual mode:
       - You type Rd as a % and it is used directly

       ### 10.4 Levering beta and WACC
       The model computes:
       - **βL = βu × (1 + (1 − Tax) × D/E)**
       - **Re = RF + βL × MRP**
       - **WACC = we×Re + wd×Rd×(1 − Tax)**

       ### 10.5 Terminal growth (g)
       This is the long-run growth rate used in terminal value.
       **Important rule:** the model needs **WACC > g** for terminal value to work.
       If WACC ≤ g, terminal value becomes invalid and sensitivity cells may be blank.
       """)

    if "guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_" not in st.session_state:
        st.session_state["guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"] else "▶ ") + "11) Step 9 — Valuation timing (date-based discounting + mid-year)",
            key="btn_guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"
    ):
        st.session_state["guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"] = not \
        st.session_state["guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"]
    if st.session_state["guide_exp_10_11__step_9___valuation_timing__date_based_discounting___mid_"]:
        st.markdown("""
       This model uses **date-based discounting**, not just “year 1, year 2”.

       You will provide:
       - **Valuation date** (today/deal date)
       - **Financial statement year-end date for first forecast year**
       - optional **mid-year convention** (subtracts 0.5 years)

       The model computes the first discount period **n₀** from dates, then:
       - discount periods = n₀, n₀+1, n₀+2, ...

       Discount factor = 1 / (1 + WACC)ⁿ
       """)

    if "guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_" not in st.session_state:
        st.session_state["guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"] else "▶ ") + "12) Step 10 — CAPEX averaging (outlier exclusions)",
            key="btn_guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"
    ):
        st.session_state["guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"] = not st.session_state[
            "guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"]
    if st.session_state["guide_exp_11_12__step_10___capex_averaging__outlier_exclusions_"]:
        st.markdown("""
       CAPEX is taken from the Cash Flow mapping.

       The model:
       - builds CAPEX history (sum of selected capex rows)
       - lets you exclude outlier years (persistently)
       - computes average CAPEX from remaining years
       - forecasts CAPEX as constant average for all forecast years

       ✅ If CAPEX is negative historically, it stays negative (cash outflow).
       """)

    if "guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_" not in st.session_state:
        st.session_state["guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"] else "▶ ") + "13) Step 11 — UFCF / FCFF calculation (core DCF engine)",
            key="btn_guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"
    ):
        st.session_state["guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"] = not st.session_state[
            "guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"]
    if st.session_state["guide_exp_12_13__step_11___ufcf___fcff_calculation__core_dcf_engine_"]:
        st.markdown("""
       Your model calculates UFCF using:

       **UFCF = EBITDA×(1 − T) + (−Depreciation×T) + ΔWC + CAPEX**

       Where:
       - EBITDA×(1−T) = after-tax operating earnings proxy
       - Depreciation×Tax adds back the tax shield (your implementation uses `-dep * tax`)
       - ΔWC is old minus new working capital
       - CAPEX is usually negative

       Then it discounts each UFCF by date-based discount factors to get PV of UFCF.
       """)

    if "guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value" not in st.session_state:
        st.session_state["guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"] else "▶ ") + "14) Step 12 — Terminal value, Enterprise value, Equity value",
            key="btn_guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"
    ):
        st.session_state["guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"] = not \
        st.session_state["guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"]
    if st.session_state["guide_exp_13_14__step_12___terminal_value__enterprise_value__equity_value"]:
        st.markdown("""
       Terminal value uses Gordon Growth:

       **TV = UFCF_last × (1 + g) / (WACC − g)**

       PV of TV = TV × DiscountFactor_last

       **Enterprise Value (EV)**  
       = Sum(PV of UFCF) + PV(TV)

       **Equity Value**  
       = EV − Net Debt

       ✅ Equity Value is saved to session state for use in Comparables and Summary pages.
       """)

    if "guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_" not in st.session_state:
        st.session_state["guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"] else "▶ ") + "15) Step 13 — Sensitivity table (WACC vs g)",
            key="btn_guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"
    ):
        st.session_state["guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"] = not st.session_state[
            "guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"]
    if st.session_state["guide_exp_14_15__step_13___sensitivity_table__wacc_vs_g_"]:
        st.markdown("""
       This grid shows how Equity Value changes when:
       - WACC changes (rows)
       - Terminal growth g changes (columns)

       You control:
       - number of WACC points
       - number of g points
       - WACC step %
       - g step %

       ⚠️ Blank cells occur when WACC ≤ g (terminal value invalid).

       The base case cell is highlighted (current WACC & g).
       Min and max are also shown.
       """)

    if "guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv" not in st.session_state:
        st.session_state["guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"] else "▶ ") + "16) Step 14 — Download FULL Excel model (formulas + sensitivity)",
            key="btn_guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"
    ):
        st.session_state["guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"] = not \
        st.session_state["guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"]
    if st.session_state["guide_exp_15_16__step_14___download_full_excel_model__formulas___sensitiv"]:
        st.markdown("""
       This button builds a **full Excel model** that includes:
       - Forecast Income Statement (with formulas)
       - Working Capital sheet
       - Inputs sheet (RF, MRP, beta, Rd, etc.)
       - DCF valuation sheet (UFCF and PV)
       - Sensitivity sheet (formula-driven grid)
       - Summary sheet

       ✅ You can edit assumptions in Excel and outputs update automatically.
       """)

    if "guide_exp_16_17__troubleshooting__common_errors___how_to_fix_" not in st.session_state:
        st.session_state["guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"] = True
    if st.button(
            ("▼ " if st.session_state[
                "guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"] else "▶ ") + "17) Troubleshooting (common errors + how to fix)",
            key="btn_guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"
    ):
        st.session_state["guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"] = not st.session_state[
            "guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"]
    if st.session_state["guide_exp_16_17__troubleshooting__common_errors___how_to_fix_"]:
        st.markdown("""
       ### “Revenue must be selected”
       You did not map the Revenue line in Income Statement mapping.

       ### “Mapping order problem”
       Your totals are mapped out of order (e.g., EBITDA appears above GP in your mapping).
       Fix by selecting the true statement line positions.

       ### “Missing FX data for statement years”
       Your FX file does not cover some statement years.
       Add FX rates for those missing years.

       ### “No FX rate found on or before closing date”
       Your chosen BS closing date has no FX rate before it in the FX file.
       Choose a later date or extend FX data.

       ### “WACC labels collided (duplicate labels)”
       Your sensitivity step is too small or decimals too low → labels become same.
       Increase decimals or increase step.

       ### “Terminal value invalid / blank sensitivity cells”
       This happens when **WACC ≤ g**.
       Reduce g or increase WACC range.
       """)


# ✅ ADD THIS inside:  with tab7:
st.markdown("### 📥 Download Quick Summary Only")

# 1) Put the Quick Summary text you want exported here (edit freely)
quick_summary_text = """
QUICK SUMMARY — Valuation App

DCF (UFCF / FCFF)
1) Upload IS + BS + CF (one Excel, 3 sheets).
   - Required sheet order: Sheet 1 = Income Statement, Sheet 2 = Balance Sheet, Sheet 3 = Cash Flows.
   - Required layout: Column A = line items (start Row 2); Row 1 (from Column B onward) = years (e.g., 2022, 2023, 2024).
2) Select currency (USD or ZWG + FX file).
   - If ZWG: upload an FX Excel with Date + Bank and/or Interbank rate columns.
3) Map: Revenue, Debt, Cash, CA, CL, Equity, Capex, Depreciation (if available).
4) Choose forecast years + revenue growth method.
5) Confirm WC% method:
   - Review historical WC% of Sales.
   - Untick “Include” to exclude outlier years from the average.
   - Choose average of included years OR most recent WC% for forecasting.
6) Enter Average Cost of Debt Zimbabwe (US$) (%).
7) Tick “Use Auto (from Excel) for RF & MRP” (or untick to override manually).
8) Select Industry / Industries (for blended βu) from the auto list or override βu manually.
9) Enter Tax rate and Terminal growth rate (g).
10) Select Valuation timing (valuation date + financial year-end date).
11) Review CAPEX History — exclude outlier years before averaging.
12) Review EV → Equity and the WACC vs g sensitivity grid.
13) Export Excel for audit trail.

DDM (Gordon Growth)
1) Enter dividend history.
2) Pick stable years for growth.
3) Confirm g and D1.
4) Set CAPM inputs (or override).
5) Check Ke > g then compute P0.
6) Enter shares for total equity value.
7) Export Excel.

Comparables (EV/EBITDA · P/B · P/E)
1) Gather all the necessary data on comparable companies and Ratios (EV/EBITDA, P/E, P/B).
3) Choose how many comparables you want to use.
2) Enter peers and also multiples.
3) Use Include flags to remove outliers (do not delete).
4) Enter Discount % to compute implied multiples.
Maintainable EBITDA/Earnings:
- Choose year range, weights, timing (from DCF).
Book Equity & Net Debt:
- Auto-pulled from DCF/BANKING where available; can override manually.

Banking (Residual Income)
1) Upload IS + BS + SoCE.
2) If ZWG: upload FX and confirm conversion.
3) Map Total Equity rows on BS.
4) Choose earnings line.
5) Choose base year.
6) Set Ke via CAPM + forecast years.
7) Enter growth assumptions + terminal g.
8) Check Ke > g then compute and export.

Summary (Blended / Weighted)
1) Run your chosen models first.
2) Select models to include.
3) Input weights (auto-normalized to 100%).
4) Enter shares and current share price.
5) Review intrinsic value + upside/downside.
6) Export Summary Excel.
""".strip()


# 2) PDF generator (quick + clean)
def build_quick_summary_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    left = 40
    top = height - 50
    line_height = 14

    y = top
    c.setFont("Times-Roman", 11)

    for line in text.split("\n"):
        if y < 60:  # new page
            c.showPage()
            c.setFont("Times-Roman", 11)
            y = top
        c.drawString(left, y, line[:120])  # keep lines safe width-wise
        y -= line_height

    c.save()
    buffer.seek(0)
    return buffer.read()


# 3) Word generator
def build_quick_summary_docx(text: str) -> bytes:
    doc = Document()
    doc.add_heading("Quick Summary — Valuation App", level=1)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # simple heading detection
        if line.isupper() or line.endswith(")"):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# 4) Download buttons
pdf_bytes = build_quick_summary_pdf(quick_summary_text)
docx_bytes = build_quick_summary_docx(quick_summary_text)

c1, c2 = st.columns(2)
with c1:
    st.download_button(
        "⬇️ Download Quick Summary (PDF)",
        data=pdf_bytes,
        file_name="quick_summary.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

with c2:
    st.download_button(
        "⬇️ Download Quick Summary (Word)",
        data=docx_bytes,
        file_name="quick_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
